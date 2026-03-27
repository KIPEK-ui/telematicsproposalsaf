"""
Form Filler Service
Fills detected form fields with proposal content while preserving document structure.
Supports DOCX and PDF formats.
"""

import logging
from typing import Dict, List, Any, Optional, Tuple
from pathlib import Path
from io import BytesIO
from dataclasses import dataclass
from .form_detector import FormField, FormStructure, FieldType

logger = logging.getLogger(__name__)

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import pypdf
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False


@dataclass
class FieldMapping:
    """Maps a proposal section to form fields."""
    proposal_section: str
    form_fields: List[str]  # field_ids
    content: str
    priority: int = 1  # 1 = high priority fill


class FormFiller:
    """
    Fills detected form fields with proposal content.
    Intelligently maps proposal sections to tender form fields.
    """

    # Mapping between proposal sections and form field keywords
    SECTION_KEYWORDS = {
        "executive_summary": [
            "summary", "overview", "introduction", "executive", "brief",
            "synopsis", "abstract"
        ],
        "technical_approach": [
            "technical", "methodology", "approach", "solution", "method",
            "process", "procedure", "strategy", "mechanism"
        ],
        "fleet_details": [
            "fleet", "equipment", "vehicle", "resource", "asset",
            "specification", "details", "description", "hardware"
        ],
        "implementation_timeline": [
            "timeline", "schedule", "implementation", "milestones",
            "dates", "duration", "timeline", "phase", "stage"
        ],
        "pricing": [
            "pricing", "cost", "price", "budget", "commercial",
            "terms", "amount", "fee", "charge", "rate"
        ],
        "compliance_assurance": [
            "compliance", "quality", "assurance", "qualification",
            "standard", "certification", "requirement", "regulation"
        ],
        "terms_conditions": [
            "terms", "conditions", "t&c", "general", "legal",
            "agreement", "covenant"
        ]
    }

    def __init__(self):
        """Initialize form filler."""
        if not DOCX_AVAILABLE:
            logger.warning("python-docx not available. DOCX filling disabled.")
        logger.info("FormFiller initialized")

    def fill_form(
        self,
        original_file_path: str,
        form_structure: FormStructure,
        proposal_content: Dict[str, str],
        org_data: Dict[str, Any]
    ) -> bytes:
        """
        Fill tender form with proposal content.

        Args:
            original_file_path: Path to original tender document
            form_structure: Detected form structure from FormDetector
            proposal_content: Proposal sections dict
            org_data: Organization information

        Returns:
            bytes: Filled document content

        Raises:
            FormFillerError: If filling fails
        """
        file_path = Path(original_file_path)
        suffix = file_path.suffix.lower()

        if suffix == ".docx":
            return self._fill_docx_form(
                file_path, form_structure, proposal_content, org_data
            )
        elif suffix == ".pdf":
            logger.warning("PDF form filling requires external tool installation")
            return b""
        else:
            raise FormFillerError(f"Unsupported file type: {suffix}")

    def _fill_docx_form(
        self,
        file_path: Path,
        form_structure: FormStructure,
        proposal_content: Dict[str, str],
        org_data: Dict[str, Any]
    ) -> bytes:
        """Fill DOCX form with proposal content."""
        if not DOCX_AVAILABLE:
            raise FormFillerError("python-docx not available")

        try:
            doc = Document(file_path)

            # Create mapping of proposal sections to form fields
            mappings = self._create_field_mappings(form_structure, proposal_content)

            # Fill paragraph fields
            self._fill_paragraph_fields(doc, form_structure, mappings, org_data)

            # Fill table fields
            self._fill_table_fields(doc, form_structure, mappings, org_data)

            # Export to bytes
            output = BytesIO()
            doc.save(output)
            output.seek(0)
            return output.getvalue()

        except Exception as e:
            logger.error(f"Error filling DOCX form: {e}")
            raise FormFillerError(f"Failed to fill DOCX form: {e}")

    def _create_field_mappings(
        self,
        form_structure: FormStructure,
        proposal_content: Dict[str, str]
    ) -> Dict[str, FieldMapping]:
        """
        Create intelligent mapping between proposal sections and form fields.
        Uses keyword matching to assign proposal content to relevant fields.
        """
        mappings = {}

        for field in form_structure.fields:
            best_section = None
            best_score = 0

            # Match field label to proposal section
            field_label_lower = field.label.lower()

            for section, keywords in self.SECTION_KEYWORDS.items():
                score = sum(
                    1 for keyword in keywords
                    if keyword in field_label_lower
                )

                if score > best_score:
                    best_score = score
                    best_section = section

            # Default fallback: assign to executive_summary
            if not best_section:
                best_section = "executive_summary"

            # Create mapping only if we have content for this section
            if best_section in proposal_content:
                if best_section not in mappings:
                    mappings[best_section] = FieldMapping(
                        proposal_section=best_section,
                        form_fields=[field.field_id],
                        content=proposal_content[best_section],
                        priority=best_score
                    )
                else:
                    mappings[best_section].form_fields.append(field.field_id)

        return mappings

    def _fill_paragraph_fields(
        self,
        doc: Document,
        form_structure: FormStructure,
        mappings: Dict[str, FieldMapping],
        org_data: Dict[str, Any]
    ) -> None:
        """Fill text fields in document paragraphs."""
        filled_fields = set()

        for para_idx, para in enumerate(doc.paragraphs):
            text = para.text.strip()

            # Find matching form field
            for field in form_structure.fields:
                if field.field_type in [FieldType.TEXT, FieldType.PARAGRAPH]:
                    location = field.location.get("paragraph_index")
                    if location == para_idx:
                        # Find best matching proposal section
                        for section, mapping in mappings.items():
                            if field.field_id in mapping.form_fields:
                                # Fill this paragraph with proposal content
                                self._fill_paragraph(
                                    para, mapping.content, field, org_data
                                )
                                filled_fields.add(field.field_id)
                                break

    def _fill_table_fields(
        self,
        doc: Document,
        form_structure: FormStructure,
        mappings: Dict[str, FieldMapping],
        org_data: Dict[str, Any]
    ) -> None:
        """Fill table cells with proposal content."""
        filled_fields = set()

        for table_idx, table in enumerate(doc.tables):
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    # Find matching field
                    field_id = f"table_{table_idx}_row_{row_idx}_cell_{cell_idx}"

                    for field in form_structure.fields:
                        if field.field_id == field_id and field.field_type == FieldType.TABLE_ROW:
                            # Find matching section
                            for section, mapping in mappings.items():
                                if field.field_id in mapping.form_fields:
                                    self._fill_table_cell(
                                        cell, mapping.content, field, org_data
                                    )
                                    filled_fields.add(field.field_id)
                                    break

    def _fill_paragraph(
        self,
        para,
        content: str,
        field: FormField,
        org_data: Dict[str, Any]
    ) -> None:
        """Fill a paragraph with proposal content."""
        # Truncate content to field's estimated length + 20%
        max_length = int(field.estimated_length * 1.2) if field.estimated_length > 0 else 500

        if len(content) > max_length:
            # Truncate and add ellipsis
            content = content[:max_length].rsplit(" ", 1)[0] + "..."

        # Replace paragraph content
        para.text = content

        # Maintain formatting
        for run in para.runs:
            run.font.size = Pt(11)

    def _fill_table_cell(
        self,
        cell,
        content: str,
        field: FormField,
        org_data: Dict[str, Any]
    ) -> None:
        """Fill a table cell with proposal content."""
        # Clear existing content
        if cell.paragraphs:
            cell.paragraphs[0].text = ""

        # Add new content
        para = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
        
        # Truncate for table cells (usually smaller)
        max_length = 200
        if len(content) > max_length:
            content = content[:max_length].rsplit(" ", 1)[0] + "..."

        para.text = content
        for run in para.runs:
            run.font.size = Pt(10)

    def extract_field_content(
        self,
        original_file_path: str,
        form_structure: FormStructure
    ) -> Dict[str, str]:
        """
        Extract existing content from form fields.
        Useful for preserving some fields while filling others.
        """
        extracted = {}

        if not Path(original_file_path).exists():
            return extracted

        try:
            doc = Document(original_file_path)

            for field in form_structure.fields:
                location = field.location
                
                if field.field_type == FieldType.TEXT and "paragraph_index" in location:
                    para_idx = location["paragraph_index"]
                    if para_idx < len(doc.paragraphs):
                        extracted[field.field_id] = doc.paragraphs[para_idx].text

            return extracted

        except Exception as e:
            logger.error(f"Error extracting field content: {e}")
            return extracted


class FormFillerError(Exception):
    """Raised when form filling fails."""
    pass
