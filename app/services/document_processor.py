"""
Document Processor Service
Converts PDF and DOCX files into text format for RAG training.
Supports PDF extraction (pdfplumber) and DOCX extraction (python-docx).
✅ ENHANCED: Adds extraction result caching to avoid re-processing
"""

import logging
import hashlib
import json
from pathlib import Path
from typing import Optional, Tuple, List, Dict, Any
from dataclasses import dataclass
from datetime import datetime

logger = logging.getLogger(__name__)


@dataclass
class ProcessedDocument:
    """Result of document processing."""
    original_path: str
    output_path: str
    document_type: str  # 'pdf', 'docx', 'txt'
    extracted_text: str
    success: bool
    error_message: Optional[str] = None


class DocumentProcessorError(Exception):
    """Raised when document processing fails."""
    pass


class DocumentProcessor:
    """
    Processes PDF and DOCX documents into text format for RAG training.
    Automatically detects file type and uses appropriate extraction method.
    ✅ ENHANCED: Caches extraction results to avoid re-processing same files
    """

    def __init__(self, output_dir: Optional[Path] = None):
        """
        Initialize document processor.
        
        Args:
            output_dir: Directory to save extracted text files.
                       Defaults to data/training_proposals/
        """
        # Convert to Path if string is passed, otherwise use default
        if output_dir:
            self.output_dir = Path(output_dir) if isinstance(output_dir, str) else output_dir
        else:
            self.output_dir = Path("data/training_proposals")
        
        self.output_dir.mkdir(parents=True, exist_ok=True)
        
        # ✅ NEW: Processing cache to track processed files
        self.cache_dir = Path("data/.processing_cache")
        self.cache_dir.mkdir(parents=True, exist_ok=True)
        self.cache_file = self.cache_dir / "extraction_cache.json"
        self._extraction_cache: Dict[str, Dict[str, Any]] = {}
        self._load_extraction_cache()
        
        logger.info(f"DocumentProcessor initialized with output: {self.output_dir}")

    def _get_file_hash(self, file_path: Path) -> str:
        """Generate hash of file for cache validation."""
        try:
            with open(file_path, 'rb') as f:
                return hashlib.md5(f.read()).hexdigest()
        except Exception:
            return ""

    def _load_extraction_cache(self) -> None:
        """Load extraction cache from disk."""
        try:
            if self.cache_file.exists():
                with open(self.cache_file, 'r', encoding='utf-8') as f:
                    self._extraction_cache = json.load(f)
                logger.debug(f"Loaded extraction cache with {len(self._extraction_cache)} entries")
        except Exception as e:
            logger.warning(f"Failed to load extraction cache: {e}")
            self._extraction_cache = {}

    def _save_extraction_cache(self) -> None:
        """Save extraction cache to disk."""
        try:
            with open(self.cache_file, 'w', encoding='utf-8') as f:
                json.dump(self._extraction_cache, f, indent=2)
            logger.debug("Saved extraction cache")
        except Exception as e:
            logger.warning(f"Failed to save extraction cache: {e}")

    def _is_cached(self, file_path: Path) -> Optional[str]:
        """Check if file was already processed and is still valid."""
        try:
            file_hash = self._get_file_hash(file_path)
            cache_key = str(file_path.absolute())
            
            if cache_key in self._extraction_cache:
                cached = self._extraction_cache[cache_key]
                if cached.get('file_hash') == file_hash and cached.get('extraction_result'):
                    logger.debug(f"Cache hit for {file_path.name}")
                    return cached['extraction_result']
        except Exception as e:
            logger.debug(f"Cache check failed: {e}")
        
        return None

    def _cache_result(self, file_path: Path, extracted_text: str) -> None:
        """Cache extraction result for future use."""
        try:
            file_hash = self._get_file_hash(file_path)
            cache_key = str(file_path.absolute())
            
            self._extraction_cache[cache_key] = {
                'file_hash': file_hash,
                'extraction_result': extracted_text,
                'cached_at': datetime.now().isoformat()
            }
            self._save_extraction_cache()
        except Exception as e:
            logger.warning(f"Failed to cache extraction result: {e}")

    def process_file(self, file_path: str) -> ProcessedDocument:
        """
        Process a single document (PDF, DOCX, or TXT).
        ✅ ENHANCED: Checks cache first to avoid re-extraction
        
        Args:
            file_path: Path to the document file
            
        Returns:
            ProcessedDocument: Processing result with extracted text
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            return ProcessedDocument(
                original_path=str(file_path),
                output_path="",
                document_type="unknown",
                extracted_text="",
                success=False,
                error_message=f"File not found: {file_path}"
            )
        
        try:
            # ✅ NEW: Check cache first
            cached_text = self._is_cached(file_path)
            if cached_text is not None:
                logger.info(f"Using cached extraction for {file_path.name}")
                
                # Still need to determine doc type and output path
                suffix = file_path.suffix.lower()
                if suffix == '.pdf':
                    doc_type = 'pdf'
                elif suffix in ['.docx', '.doc']:
                    doc_type = 'docx'
                elif suffix == '.txt':
                    doc_type = 'txt'
                else:
                    doc_type = 'unknown'
                
                # Save to output if not already there
                output_file = self._save_text(file_path, cached_text, doc_type)
                
                return ProcessedDocument(
                    original_path=str(file_path),
                    output_path=str(output_file),
                    document_type=doc_type,
                    extracted_text=cached_text,
                    success=True
                )
            
            # Cache miss - process normally
            # Detect document type
            suffix = file_path.suffix.lower()
            
            if suffix == '.pdf':
                text = self._extract_pdf(file_path)
                doc_type = 'pdf'
            elif suffix in ['.docx', '.doc']:
                text = self._extract_docx(file_path)
                doc_type = 'docx'
            elif suffix == '.txt':
                text = self._extract_txt(file_path)
                doc_type = 'txt'
            else:
                raise DocumentProcessorError(f"Unsupported file type: {suffix}")
            
            # ✅ NEW: Cache the extraction result
            self._cache_result(file_path, text)
            
            # Save to text file
            output_file = self._save_text(file_path, text, doc_type)
            
            logger.info(f"Processed {doc_type.upper()} file: {file_path.name} → {output_file.name}")
            
            return ProcessedDocument(
                original_path=str(file_path),
                output_path=str(output_file),
                document_type=doc_type,
                extracted_text=text,
                success=True
            )
            
        except Exception as e:
            logger.error(f"Failed to process {file_path}: {str(e)}")
            return ProcessedDocument(
                original_path=str(file_path),
                output_path="",
                document_type="unknown",
                extracted_text="",
                success=False,
                error_message=str(e)
            )

    def process_directory(self, directory: str) -> List[ProcessedDocument]:
        """
        Process all documents in a directory.
        Recursively searches for PDF, DOCX, and TXT files.
        
        Args:
            directory: Directory containing documents
            
        Returns:
            List[ProcessedDocument]: Results for all processed documents
        """
        dir_path = Path(directory)
        
        if not dir_path.exists():
            raise DocumentProcessorError(f"Directory not found: {directory}")
        
        results = []
        
        # Find all supported document types
        file_patterns = ['*.pdf', '*.docx', '*.doc', '*.txt']
        files = []
        for pattern in file_patterns:
            files.extend(dir_path.rglob(pattern))
        
        logger.info(f"Found {len(files)} documents in {directory}")
        
        for file_path in files:
            result = self.process_file(str(file_path))
            results.append(result)
        
        # Summary
        successful = sum(1 for r in results if r.success)
        logger.info(f"Processing complete: {successful}/{len(results)} successful")
        
        return results

    # ========================
    # Extraction Methods
    # ========================

    def _extract_pdf(self, file_path: Path) -> str:
        """
        Extract text from PDF file using pdfplumber.
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            str: Extracted text from all pages
            
        Raises:
            DocumentProcessorError: If extraction fails
        """
        try:
            import pdfplumber
        except ImportError:
            raise DocumentProcessorError(
                "pdfplumber not installed. Run: pip install pdfplumber"
            )
        
        try:
            text_parts = []
            with pdfplumber.open(file_path) as pdf:
                logger.info(f"Extracting text from {len(pdf.pages)} pages")
                
                for i, page in enumerate(pdf.pages):
                    try:
                        text = page.extract_text()
                        if text:
                            text_parts.append(f"--- Page {i+1} ---\n{text}")
                    except Exception as page_error:
                        logger.warning(f"Failed to extract page {i+1}: {page_error}")
                        continue
            
            full_text = "\n\n".join(text_parts)
            
            if not full_text.strip():
                raise DocumentProcessorError("No text extracted from PDF (possibly empty or scanned)")
            
            return full_text
            
        except Exception as e:
            raise DocumentProcessorError(f"PDF extraction failed: {str(e)}")

    def _extract_docx(self, file_path: Path) -> str:
        """
        Extract text from DOCX file using python-docx.
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            str: Extracted text from all paragraphs
            
        Raises:
            DocumentProcessorError: If extraction fails
        """
        try:
            from docx import Document
        except ImportError:
            raise DocumentProcessorError(
                "python-docx not installed. Run: pip install python-docx"
            )
        
        try:
            doc = Document(file_path)
            text_parts = []
            
            # Extract from paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # Extract from tables
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        row_text.append(cell.text.strip())
                    table_text.append(" | ".join(row_text))
                if table_text:
                    text_parts.append("\n".join(table_text))
            
            full_text = "\n\n".join(text_parts)
            
            if not full_text.strip():
                raise DocumentProcessorError("No text extracted from DOCX (possibly empty)")
            
            return full_text
            
        except Exception as e:
            raise DocumentProcessorError(f"DOCX extraction failed: {str(e)}")

    def _extract_txt(self, file_path: Path) -> str:
        """
        Extract text from TXT file (simple read).
        
        Args:
            file_path: Path to TXT file
            
        Returns:
            str: File contents
            
        Raises:
            DocumentProcessorError: If reading fails
        """
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
            
            if not text.strip():
                raise DocumentProcessorError("TXT file is empty")
            
            return text
            
        except Exception as e:
            raise DocumentProcessorError(f"TXT read failed: {str(e)}")

    # ========================
    # File Output
    # ========================

    def _save_text(self, original_path: Path, text: str, doc_type: str) -> Path:
        """
        Save extracted text to output directory.
        
        Args:
            original_path: Original document path
            text: Extracted text
            doc_type: Document type (pdf, docx, txt)
            
        Returns:
            Path: Path to saved text file
        """
        # Create output filename from original filename
        output_name = f"{original_path.stem}_{doc_type}.txt"
        output_path = self.output_dir / output_name
        
        try:
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(text)
            
            logger.info(f"Saved extracted text to: {output_path}")
            return output_path
            
        except Exception as e:
            raise DocumentProcessorError(f"Failed to save text file: {str(e)}")


# ========================
# Utility Functions
# ========================

def process_training_documents(source_dir: str, output_dir: Optional[str] = None) -> Tuple[int, int, List[str]]:
    """
    Batch process all documents in a directory for RAG training.
    
    Args:
        source_dir: Directory containing PDF, DOCX, TXT files
        output_dir: Directory to save extracted text (defaults to data/training_proposals/)
        
    Returns:
        Tuple of (successful_count, total_count, error_messages)
    """
    processor = DocumentProcessor(Path(output_dir) if output_dir else None)
    results = processor.process_directory(source_dir)
    
    errors = [r.error_message for r in results if not r.success and r.error_message]
    successful = sum(1 for r in results if r.success)
    
    return successful, len(results), errors


def get_document_processor() -> DocumentProcessor:
    """Get or create document processor instance."""
    return DocumentProcessor()
