"""
Tender Parser Service
Parses and extracts structured data from tenders in multiple formats (PDF, text, form).
Production-ready with validation, error handling, and structured output.
Enhanced with pypdf2, pdfplumber table extraction, and deepseek-ocr integration.
"""

import re
import logging
import json
import io
from typing import Dict, Any, List, Optional, Tuple
from dataclasses import dataclass, field
from enum import Enum
from abc import ABC, abstractmethod

# Try to import PDF libraries
try:
    import pypdf
    PYPDF_AVAILABLE = True
except ImportError:
    PYPDF_AVAILABLE = False

try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False

try:
    from paddleocr import PaddleOCR
    PADDLE_OCR_AVAILABLE = True
except ImportError:
    PADDLE_OCR_AVAILABLE = False

try:
    from pdf2image import convert_from_bytes
    PDF2IMAGE_AVAILABLE = True
except ImportError:
    PDF2IMAGE_AVAILABLE = False

try:
    import pytesseract
    PYTESSERACT_AVAILABLE = True
except ImportError:
    PYTESSERACT_AVAILABLE = False

try:
    from docx import Document
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False

logger = logging.getLogger(__name__)


class TenderSourceType(Enum):
    """Supported tender input types."""
    PDF = "pdf"
    TEXT = "text"
    FORM = "form"


@dataclass
class TenderDocument:
    """
    Standardized representation of a parsed tender.
    
    Attributes:
        source_type: Type of input (pdf, text, form)
        title: Extracted or provided tender title
        raw_content: Original unprocessed content
        sections: Extracted sections with structure
        technical_requirements: Parsed technical specs
        fleet_details: Fleet/equipment requirements
        timeline: Project timeline/deadline
        budget_info: Budget/pricing constraints
        tender_no: Tender reference number
        address: Tender issuer address
        email: Tender issuer email
        phone_number: Tender issuer phone number
        bid_validity: Bid validity period
        metadata: Additional metadata from parsing
    """
    source_type: TenderSourceType
    title: str
    raw_content: str
    sections: Dict[str, str] = field(default_factory=dict)
    technical_requirements: Dict[str, Any] = field(default_factory=dict)
    fleet_details: Dict[str, Any] = field(default_factory=dict)
    timeline: Optional[str] = None
    budget_info: Optional[str] = None
    tender_no: Optional[str] = None
    address: Optional[str] = None
    email: Optional[str] = None
    phone_number: Optional[str] = None
    bid_validity: Optional[str] = None
    metadata: Dict[str, Any] = field(default_factory=dict)

    def to_dict(self) -> Dict[str, Any]:
        """Convert to dictionary for database storage."""
        return {
            'source_type': self.source_type.value,
            'title': self.title,
            'raw_content': self.raw_content,
            'sections': self.sections,
            'technical_requirements': self.technical_requirements,
            'fleet_details': self.fleet_details,
            'timeline': self.timeline,
            'budget_info': self.budget_info,
            'tender_no': self.tender_no,
            'address': self.address,
            'email': self.email,
            'bid_validity': self.bid_validity,
            'metadata': self.metadata
        }


class TenderParsingError(Exception):
    """Raised when tender parsing fails."""
    pass


class BaseTenderParser(ABC):
    """
    Abstract base class for tender parsers.
    Defines interface for different input format parsers.
    """

    @abstractmethod
    def parse(self) -> TenderDocument:
        """Parse tender and return standardized document."""
        pass

    @staticmethod
    def _extract_sections(text: str) -> Dict[str, str]:
        """
        Extract common sections from unstructured text.
        Looks for common patterns: Requirements, Overview, Timeline, etc.
        """
        sections = {}
        
        # Common section headers (case-insensitive)
        section_patterns = [
            ('overview', r'overview|background|introduction'),
            ('requirements', r'requirements|specifications|technical specs'),
            ('deliverables', r'deliverables|execution|scope'),
            ('timeline', r'timeline|schedule|deadline'),
            ('budget', r'budget|pricing|cost|investment'),
            ('evaluation', r'evaluation criteria|selection criteria'),
            ('terms', r'terms and conditions|terms & conditions|t&c'),
        ]
        
        for section_name, pattern in section_patterns:
            # Find section header
            match = re.search(rf'(?i)^.*{pattern}.*?:\s*\n(.*?)(?=\n(?:overview|requirements|deliverables|timeline|budget|evaluation|terms|$))',
                            text, re.MULTILINE | re.DOTALL)
            if match and match.group(1):
                sections[section_name] = match.group(1).strip()
        
        return sections

    @staticmethod
    def _extract_technical_requirements(text: str) -> Dict[str, Any]:
        """Extract technical requirements and specifications."""
        requirements = {
            'specifications': [],
            'compliance': [],
            'performance': []
        }
        
        # Look for common technical patterns
        # Performance metrics
        performance_patterns = [
            r'(?i)performance.*?[:=]\s*(.+?)(?:\n|$)',
            r'(?i)throughput.*?[:=]\s*(.+?)(?:\n|$)',
            r'(?i)latency.*?[:=]\s*(.+?)(?:\n|$)',
            r'(?i)uptime.*?[:=]\s*(.+?)(?:\n|$)',
        ]
        
        for pattern in performance_patterns:
            matches = re.findall(pattern, text)
            requirements['performance'].extend(matches)
        
        # Compliance requirements
        compliance_patterns = [
            r'(?i)iso\s*[\d\-]+',
            r'(?i)compliance.*?[:=]\s*(.+?)(?:\n|$)',
            r'(?i)certifications?.*?[:=]\s*(.+?)(?:\n|$)',
        ]
        
        for pattern in compliance_patterns:
            matches = re.findall(pattern, text)
            requirements['compliance'].extend(matches)
        
        return requirements

    @staticmethod
    def _extract_fleet_details(text: str) -> Dict[str, Any]:
        """Extract fleet and equipment requirements."""
        fleet = {
            'vehicles': [],
            'equipment': [],
            'specifications': []
        }
        
        # Look for vehicle/fleet mentions
        vehicle_patterns = [
            r'(?i)(?:truck|vehicle|car|van|bus)\s+(?:models?|types?|quantities?)?[:=]?\s*(.+?)(?:\n|$)',
            r'(?i)fleet.*?[:=]\s*(.+?)(?:\n|$)',
            r'(?i)(?:number|quantity|count).*?(?:truck|vehicle).*?[:=]?\s*(.+?)(?:\n|$)',
        ]
        
        for pattern in vehicle_patterns:
            matches = re.findall(pattern, text)
            fleet['vehicles'].extend(matches)
        
        # Equipment specifications
        equipment_patterns = [
            r'(?i)equipment.*?[:=]\s*(.+?)(?:\n|$)',
            r'(?i)gps|tracking|communication|safety.*?[:=]?\s*(.+?)(?:\n|$)',
        ]
        
        for pattern in equipment_patterns:
            matches = re.findall(pattern, text)
            fleet['equipment'].extend(matches)
        
        return fleet

    @staticmethod
    def _extract_timeline(text: str) -> Optional[str]:
        """Extract timeline/deadline information including time components."""
        # Look for lines with specific date/time indicators
        lines = text.split('\n')
        
        for i, line in enumerate(lines[:100]):  # Search first 100 lines
            line_lower = line.lower()
            
            # Must contain specific keywords
            has_key_phrase = any(phrase in line_lower for phrase in [
                'deadline', 'closing', 'submission', 'closing date', 
                'submission deadline', 'date for', 'last date'
            ])
            
            if not has_key_phrase:
                continue
            
            # Must contain date indicators
            has_date_indicator = any(indicator in line for indicator in [
                '202', '20th', '21st', '22nd', '23rd', '24th', '25th', '26th', '27th', '28th', '29th', '30th', '31st',
                'Jan', 'Feb', 'Mar', 'Apr', 'May', 'Jun', 'Jul', 'Aug', 'Sep', 'Oct', 'Nov', 'Dec',
                'days', 'weeks', 'months', 'hours'
            ]) or any(char.isdigit() for char in line[-20:])
            
            if has_date_indicator:
                # Start with the current line
                result = line.strip()
                
                # If time is not in current line, check next line for time patterns
                if not re.search(r'at\s+\d{1,2}:\d{2}', result, re.IGNORECASE):
                    # Check next line for time information
                    if i + 1 < len(lines):
                        next_line = lines[i + 1].strip()
                        # If next line contains time, append it
                        if re.search(r'\d{1,2}:\d{2}\s*(?:a\.?m|p\.?m|am|pm|hours)', next_line, re.IGNORECASE):
                            result = result + ' ' + next_line
                
                # Clean up dots and page numbers
                result = re.sub(r'\s*\.+\s*\d+\s*', '', result)  # Remove "........ 10"
                result = re.sub(r'(.{5,300})[\s.]*$', r'\1', result)  # Trim to reasonable length (increased from 150)
                if 10 < len(result) < 350:
                    return result
        
        return None

    @staticmethod
    def _extract_budget(text: str) -> Optional[str]:
        """Extract budget/pricing information - conservative approach."""
        lines = text.split('\n')
        
        for i, line in enumerate(lines[:100]):  # Search first 100 lines
            line_lower = line.lower()
            
            # Must contain specific budget keywords
            has_budget_keyword = any(keyword in line_lower for keyword in [
                'budget', 'estimated cost', 'tender value', 'contract value', 
                'estimated amount', 'price range', 'cost estimate'
            ])
            
            if not has_budget_keyword:
                continue
            
            # Must have numbers (currency indicators)
            has_number = any(indicator in line for indicator in [
                '$', '€', '£', 'ksh', 'kshs', 'kenyan shillings', 'KES'
            ]) or bool(re.search(r'\d{3,}', line))  # At least 3-digit number
            
            if has_number:
                result = line.strip()
                # Remove excess dots and formatting
                result = re.sub(r'\.{3,}', '', result)
                result = re.sub(r'\s+', ' ', result)
                if 10 < len(result) < 250:
                    return result
        
        return None

    @staticmethod
    def _extract_tender_no(text: str) -> Optional[str]:
        """Extract tender/reference number - handles various formats."""
        lines = text.split('\n')
        
        for line in lines[:80]:  # Increased search range
            line_lower = line.lower()
            line_normalized = re.sub(r'\.', '', line_lower)
            
            # Pattern 1: Explicit keyword-based extraction (e.g., "TENDER NO. KRA/HQS/DP-016/2025-2026")
            if any(keyword in line_normalized for keyword in ['tender no', 'ref no', 'reference', 'tender ref', 'bid no', 'tender number']):
                # Try multiple regex patterns to capture the number
                patterns = [
                    r'(?:tender\s+no\.?|ref\s+no\.?|reference|tender\s+ref|bid\s+no\.?|tender\s+number)\s*[:=]?\s*([A-Z0-9\-/\.]+)',  # Explicit keyword
                    r'(?:tender\s+no\.?|ref\s+no\.?)\s*[:=]?\s*([A-Z0-9\-/\.]+)',  # Simplified
                ]
                
                for pattern in patterns:
                    match = re.search(pattern, line, re.IGNORECASE)
                    if match:
                        result = match.group(1).strip().rstrip('.')
                        if len(result) > 3:  # Ensure it's a meaningful tender number
                            return result
            
            # Pattern 2: Fallback - look for tender/reference patterns even without explicit keyword
            # Matches: KRA/HQS/DP-016/2025-2026 or similar patterns
            if any(keyword in line_normalized for keyword in ['kra', 'tenderno', 'ref', 'tenderid', 'tender id']):
                # Look for alphanumeric with slashes/dashes pattern (common in tender numbers)
                match = re.search(r'\b([A-Z][A-Z0-9]*(?:[/-][A-Z0-9]+){2,})\b', line)
                if match:
                    result = match.group(1).strip()
                    if len(result) > 5:  # Ensure meaningful length
                        return result
        
        return None

    @staticmethod
    def _extract_address(text: str) -> Optional[str]:
        """Extract tender issuer address."""
        lines = text.split('\n')
        
        # First, try to find address by keywords
        for i, line in enumerate(lines[:80]):
            line_lower = line.lower()
            
            # Normalize dots for keyword detection (P.O. BOX -> pobox)
            line_normalized = re.sub(r'\.', '', line_lower)
            
            # Look for address keywords
            if any(keyword in line_normalized for keyword in ['address', 'location', 'pobox', 'office', 'headquarters', 'times tower']):
                # Get this line and potentially combine with next 2-3 lines if they look like address continuation
                result_lines = [line.strip()]
                
                for j in range(i+1, min(i+4, len(lines))):
                    next_line = lines[j].strip()
                    # Stop if we hit a keyword line or metadata
                    if (next_line and not next_line[0].isdigit() and 
                        not any(kw in next_line.lower() for kw in ['tender', 'email', 'phone', 'date', 'closing'])):
                        result_lines.append(next_line)
                    else:
                        break
                
                result = '\n'.join(result_lines)
                if len(result) > 10:
                    return result.strip()
        
        # Fallback: Look for address between tender number and email
        email_match = None
        email_line_idx = -1
        tender_line_idx = -1
        
        for i, line in enumerate(lines[:100]):
            line_lower = line.lower()
            if '@' in line and 'email' in line_lower.lower():
                email_line_idx = i
                break
            if re.search(r'[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}', line):
                email_line_idx = i
                break
        
        for i, line in enumerate(lines[:50]):
            if 'tender' in line.lower() and 'no' in line.lower():
                tender_line_idx = i
                break
        
        # If we found both tender and email positions, extract lines between them as address
        if tender_line_idx >= 0 and email_line_idx > tender_line_idx:
            address_lines = []
            for i in range(tender_line_idx + 1, email_line_idx):
                line = lines[i].strip()
                if line and not any(keyword in line.lower() for keyword in ['closing', 'date']):
                    address_lines.append(line)
            
            if address_lines:
                result = '\n'.join(address_lines)
                if len(result) > 10:
                    return result.strip()
        
        return None


    @staticmethod
    def _extract_email(text: str) -> Optional[str]:
        """Extract contact email addresses."""
        # Look for email pattern
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        matches = re.findall(email_pattern, text[:4000])  # Search first 4000 chars
        
        if matches:
            # Return first email, or combine multiple if found
            return matches[0]
        
        # Fallback: look for "email:" followed by text
        lines = text.split('\n')
        for line in lines[:60]:
            if 'email' in line.lower():
                # Extract after "email:" or "email ="
                match = re.search(r'email[:\s=]+(.+?)(?:\n|$)', line, re.IGNORECASE)
                if match:
                    email_text = match.group(1).strip()
                    if email_text:
                        return email_text
        
        return None

    @staticmethod
    def _extract_bid_validity(text: str) -> Optional[str]:
        """Extract bid validity period."""
        lines = text.split('\n')
        
        for line in lines[:100]:
            line_lower = line.lower()
            
            # Look for bid validity keywords
            if any(keyword in line_lower for keyword in ['bid validity', 'validity period', 'bid valid', 'valid for', 'offer validity']):
                result = line.strip()
                # Remove excessive dots and formatting
                result = re.sub(r'\.{3,}', '', result)
                result = re.sub(r'\s+', ' ', result)
                if 10 < len(result) < 200:
                    return result
        
        return None

    @staticmethod
    def _extract_phone_number(text: str) -> Optional[str]:
        """Extract phone number(s) from tender document."""
        # Phone number patterns
        phone_patterns = [
            r'(?:TEL|PHONE|CONTACT)[:\s]+(\+?\d[\d\s\-\(\)]+\d)',  # TEL: +254 02 281 7022
            r'(?:TEL|PHONE)[:\s=]+([+\d\s\-\(\)]+)',  # With more flexible spacing
            r'\+254[\d\s\-\(\)]+\d',  # Kenya direct: +254...
            r'0[1-9][\d\s\-]{5,}',  # Local Kenya: 02 281 7022
        ]
        
        for pattern in phone_patterns:
            match = re.search(pattern, text[:3000], re.IGNORECASE)
            if match:
                phone = match.group(1) if match.groups() else match.group(0)
                phone = phone.strip()
                if phone and len(phone) > 5:
                    return phone
        
        return None


class PDFTenderParser(BaseTenderParser):
    """
    Enhanced PDF tender parser using:
    - pypdf: PDF text extraction (primary)
    - pdfplumber: Table extraction
    - PaddleOCR: OCR for scanned PDFs (fallback)
    """

    def __init__(self, file_bytes: bytes):
        """
        Initialize PDF parser.
        
        Args:
            file_bytes: PDF file content as bytes
            
        Raises:
            TenderParsingError: If no PDF libraries available
        """
        if not PYPDF_AVAILABLE and not PDFPLUMBER_AVAILABLE:
            raise TenderParsingError(
                "PDF parsing not available. Install: pip install pypdf pdfplumber paddleocr pytesseract pdf2image pillow"
            )
        self.file_bytes = file_bytes
        self.pdf_file = io.BytesIO(file_bytes)

    def parse(self) -> TenderDocument:
        """
        Parse PDF using multi-method extraction:
        1. Text extraction with pypdf
        2. Table extraction with pdfplumber
        3. OCR fallback with PaddleOCR for scanned PDFs
        
        Returns:
            TenderDocument: Parsed tender
            
        Raises:
            TenderParsingError: If PDF cannot be parsed
        """
        try:
            text_content = ""
            tables_content = ""
            ocr_content = ""
            page_count = 0
            extraction_methods = []
            
            # Method 1: Extract text using pypdf
            if PYPDF_AVAILABLE:
                logger.info("Extracting text with pypdf...")
                self.pdf_file.seek(0)
                try:
                    from pypdf import PdfReader
                    reader = PdfReader(self.pdf_file)
                    page_count = len(reader.pages)
                    
                    for page_num, page in enumerate(reader.pages):
                        page_text = page.extract_text()
                        if page_text:
                            text_content += f"\n--- Page {page_num + 1} ---\n{page_text}"
                    
                    if text_content.strip():
                        extraction_methods.append('pypdf')
                except Exception as e:
                    logger.warning(f"pypdf extraction failed: {str(e)}")
            
            # Fallback to pdfplumber if pypdf didn't work
            if not text_content.strip() and PDFPLUMBER_AVAILABLE:
                logger.info("Fallback: Extracting text with pdfplumber...")
                self.pdf_file.seek(0)
                try:
                    with pdfplumber.open(self.pdf_file) as pdf:
                        page_count = len(pdf.pages)
                        for page_num, page in enumerate(pdf.pages):
                            page_text = page.extract_text()
                            if page_text:
                                text_content += f"\n--- Page {page_num + 1} ---\n{page_text}"
                        
                        if text_content.strip():
                            extraction_methods.append('pdfplumber_text')
                except Exception as e:
                    logger.warning(f"pdfplumber text extraction failed: {str(e)}")
            
            # Method 2: Extract tables using pdfplumber
            if PDFPLUMBER_AVAILABLE and text_content.strip():
                logger.info("Extracting tables with pdfplumber...")
                self.pdf_file.seek(0)
                try:
                    with pdfplumber.open(self.pdf_file) as pdf:
                        for page_num, page in enumerate(pdf.pages):
                            tables = page.extract_tables()
                            if tables:
                                tables_content += f"\n--- Tables on Page {page_num + 1} ---\n"
                                for table_idx, table in enumerate(tables):
                                    tables_content += f"Table {table_idx + 1}:\n"
                                    for row in table:
                                        tables_content += " | ".join([str(cell) if cell else "" for cell in row])
                                        tables_content += "\n"
                                    tables_content += "\n"
                        
                        if tables_content.strip():
                            extraction_methods.append('pdfplumber_tables')
                except Exception as e:
                    logger.warning(f"pdfplumber table extraction failed: {str(e)}")
            
            # Method 3: Use OCR if text extraction was poor
            if text_content.strip():
                extracted_text_quality = len(text_content.split()) / max(1, page_count)
            else:
                extracted_text_quality = 0
            
            if extracted_text_quality < 50 and PADDLE_OCR_AVAILABLE:
                logger.info("Text extraction poor, attempting PaddleOCR...")
                try:
                    ocr_content = self._extract_with_ocr()
                    if ocr_content.strip():
                        extraction_methods.append('ocr')
                except Exception as e:
                    logger.warning(f"OCR extraction failed: {str(e)}")
            
            # Combine all extracted content
            full_text = text_content
            if tables_content:
                full_text += "\n" + tables_content
            if ocr_content:
                full_text += "\n[OCR Content]\n" + ocr_content
            
            if not full_text.strip():
                raise TenderParsingError("Could not extract any text from PDF")
            
            # Extract title from first page
            first_page_text = text_content.split('\n')[0:100]
            first_page_str = '\n'.join(first_page_text)
            title = self._extract_title_from_text(first_page_str)
            
            logger.info(f"PDF parsed successfully. Title: {title}")
            logger.info(f"Extraction methods used: {', '.join(extraction_methods)}")
            
            return TenderDocument(
                source_type=TenderSourceType.PDF,
                title=title,
                raw_content=full_text,
                sections=self._extract_sections(full_text),
                technical_requirements=self._extract_technical_requirements(full_text),
                fleet_details=self._extract_fleet_details(full_text),
                timeline=self._extract_timeline(full_text),
                budget_info=self._extract_budget(full_text),
                tender_no=self._extract_tender_no(full_text),
                address=self._extract_address(full_text),
                email=self._extract_email(full_text),
                phone_number=self._extract_phone_number(full_text),
                bid_validity=self._extract_bid_validity(full_text),
                metadata={
                    'page_count': page_count,
                    'extraction_methods': extraction_methods
                }
            )
        
        except TenderParsingError:
            raise
        except Exception as e:
            logger.error(f"PDF parsing failed: {str(e)}")
            raise TenderParsingError(f"Failed to parse PDF: {str(e)}")

    def _extract_with_ocr(self) -> str:
        """
        Extract text from PDF using OCR (for scanned PDFs).
        Tries pytesseract first (faster), then falls back to PaddleOCR.
        Uses pdf2image to convert PDF pages to images for OCR.
        
        Returns:
            str: Extracted text from OCR
        """
        if not PDF2IMAGE_AVAILABLE:
            logger.warning("pdf2image not available for OCR processing")
            return ""
        
        if not PYTESSERACT_AVAILABLE and not PADDLE_OCR_AVAILABLE:
            logger.warning("Neither pytesseract nor PaddleOCR available for OCR processing")
            logger.info("Install with: pip install pytesseract paddleocr pillow pdf2image")
            return ""
        
        try:
            from pdf2image import convert_from_bytes
            
            # Convert PDF pages to images
            logger.info("Converting PDF pages to images for OCR...")
            images = convert_from_bytes(self.file_bytes, dpi=150)
            
            ocr_text = ""
            max_pages = min(10, len(images))  # Limit to first 10 pages for speed
            
            # Try pytesseract first (faster)
            if PYTESSERACT_AVAILABLE:
                logger.info("Attempting OCR with pytesseract...")
                try:
                    for page_num, image in enumerate(images[:max_pages]):
                        logger.info(f"Running pytesseract on page {page_num + 1}/{max_pages}...")
                        try:
                            page_text = pytesseract.image_to_string(image)
                            if page_text.strip():
                                ocr_text += f"\n--- Pytesseract Page {page_num + 1} ---\n{page_text}"
                        except Exception as e:
                            logger.warning(f"Pytesseract failed on page {page_num + 1}: {str(e)}")
                    
                    if ocr_text.strip():
                        logger.info(f"Pytesseract extracted {len(ocr_text.split())} words")
                        return ocr_text
                except Exception as e:
                    logger.warning(f"Pytesseract OCR failed: {str(e)}")
            
            # Fallback to PaddleOCR
            if PADDLE_OCR_AVAILABLE:
                logger.info("Fallback: Attempting OCR with PaddleOCR...")
                try:
                    from paddleocr import PaddleOCR
                    
                    logger.info("Initializing PaddleOCR...")
                    ocr = PaddleOCR(use_angle_cls=True, lang='en')
                    
                    ocr_text = ""
                    for page_num, image in enumerate(images[:max_pages]):
                        logger.info(f"Running PaddleOCR on page {page_num + 1}/{max_pages}...")
                        try:
                            result = ocr.ocr(image, cls=True)
                            
                            # Extract text from OCR result
                            page_text = ""
                            for line in result:
                                if line:
                                    for cell in line:
                                        text = cell[0]
                                        confidence = cell[1]
                                        if confidence > 0.5:  # Confidence > 50%
                                            page_text += text + " "
                            
                            if page_text.strip():
                                ocr_text += f"\n--- PaddleOCR Page {page_num + 1} ---\n{page_text}"
                        except Exception as e:
                            logger.warning(f"PaddleOCR failed on page {page_num + 1}: {str(e)}")
                    
                    if ocr_text.strip():
                        logger.info(f"PaddleOCR extracted {len(ocr_text.split())} words")
                        return ocr_text
                except Exception as e:
                    logger.warning(f"PaddleOCR extraction failed: {str(e)}")
            
            return ocr_text
        
        except ImportError as e:
            logger.warning(f"OCR dependencies not installed: {str(e)}")
            logger.info("Install with: pip install pytesseract paddleocr pdf2image pillow")
            return ""
        except Exception as e:
            logger.warning(f"OCR extraction error: {str(e)}")
            return ""

    @staticmethod
    def _extract_title_from_text(text: str) -> str:
        """Extract tender title from text - handles multi-line titles."""
        lines = [l.strip() for l in text.split('\n') if l.strip()]
        
        # Strategy 1: Look for all-caps lines, but check if they're complete
        for idx, line in enumerate(lines[:20]):
            if (len(line) > 15 and len(line) < 300 and
                line.isupper() and 
                not all(c in '.-=_#*' for c in line) and
                line.count(' ') >= 2):
                
                # Check if this line looks incomplete (ends with incomplete word/phrase)
                incomplete_endings = (' FOR', ' OF', ' AND', ' WITH', ' TO', ' THE', ' A', ' AN', 'FOR A')
                
                if line.endswith(incomplete_endings) and idx + 1 < len(lines):
                    next_line = lines[idx + 1].strip()
                    # If next line is also all-caps and not a reference number, join them
                    if (next_line and next_line.isupper() and 
                        not next_line.startswith(('REF', 'TENDER', 'NO.', 'RFQ', 'RFP')) and
                        len(line) + len(next_line) < 350):
                        combined = f"{line} {next_line}"
                        # Stop if we hit a line that looks like metadata
                        if not combined.endswith((':', 'YEARS', 'MONTHS', 'DAYS')):
                            return combined
                        # If it ends with YEARS/valid ending, return it
                        if combined.endswith(('YEARS', 'MONTHS', 'DAYS', ')', 'SERVICES', 'PRODUCTS')):
                            return combined
                
                # Line looks complete on its own
                if not line.endswith(incomplete_endings):
                    return line
        
        # Strategy 2: Join first 2-3 meaningful lines if they form a compound title
        candidate_lines = []
        for line in lines[:10]:
            if (20 <= len(line) <= 200 and 
                line.count(' ') >= 1 and
                not line.endswith(':') and
                not line.startswith('---') and
                not any(c in line for c in ['[', ']', '()']) and
                not any(x in line for x in ['TENDER', 'KRA/', 'TEL:', 'EMAIL:', 'P.O.'])):
                candidate_lines.append(line)
            if len(candidate_lines) >= 2:
                break
        
        if len(candidate_lines) >= 2:
            combined = " ".join(candidate_lines[:2])
            if 20 < len(combined) < 350:
                return combined
        
        if candidate_lines and len(candidate_lines[0]) > 15:
            return candidate_lines[0]
        
        # Fallback: first substantial line
        for line in lines:
            if 15 <= len(line) <= 300:
                return line
        
        return "Untitled Tender"


class TextTenderParser(BaseTenderParser):
    """Parse tender from plain text or email input."""

    def __init__(self, text_content: str, title: Optional[str] = None):
        """
        Initialize text parser.
        
        Args:
            text_content: Tender content as text
            title: Optional title to override extraction
        """
        self.text_content = text_content
        self.provided_title = title

    def parse(self) -> TenderDocument:
        """
        Parse text tender and extract structured information.
        
        Returns:
            TenderDocument: Parsed tender
            
        Raises:
            TenderParsingError: If text is empty or invalid
        """
        if not self.text_content or not self.text_content.strip():
            raise TenderParsingError("Text content cannot be empty")
        
        # Extract title
        title = self.provided_title or self._extract_title_from_text(self.text_content)
        
        logger.info(f"Text tender parsed. Title: {title}")
        
        return TenderDocument(
            source_type=TenderSourceType.TEXT,
            title=title,
            raw_content=self.text_content,
            sections=self._extract_sections(self.text_content),
            technical_requirements=self._extract_technical_requirements(self.text_content),
            fleet_details=self._extract_fleet_details(self.text_content),
            timeline=self._extract_timeline(self.text_content),
            budget_info=self._extract_budget(self.text_content),
            tender_no=self._extract_tender_no(self.text_content),
            address=self._extract_address(self.text_content),
            email=self._extract_email(self.text_content),
            phone_number=self._extract_phone_number(self.text_content),
            bid_validity=self._extract_bid_validity(self.text_content),
            metadata={'input_lines': len(self.text_content.split('\n'))}
        )

    @staticmethod
    def _extract_title_from_text(text: str) -> str:
        """Extract title from text - same as PDFTenderParser for consistency."""
        lines = text.split('\n')
        
        # Strategy 1: Look for all-caps lines first (titles are often all caps)
        for line in lines[:20]:
            line = line.strip()
            if (len(line) > 15 and len(line) < 300 and
                line.isupper() and 
                not all(c in '.-=_#*' for c in line) and
                line.count(' ') >= 2):  # Must have at least 2 words
                return line
        
        # Strategy 2: First meaningful line that looks like a title (substantive, reasonable length)
        best_title = ""
        for line in lines[:15]:
            line = line.strip()
            # Title should be: 20-300 chars, have multiple words, not end with colon
            if (20 <= len(line) <= 300 and 
                line.count(' ') >= 2 and
                not line.endswith(':') and
                not line.startswith('---') and
                not any(c in line for c in ['[', ']', '()', '{}', 'http', 'email'])):
                # Prefer longer lines (more likely to be full title)
                if len(line) > len(best_title):
                    best_title = line
        
        if best_title and len(best_title) > 15:
            return best_title
        
        # Strategy 3: Join first 2-3 lines if they form a complete title
        first_lines = [l.strip() for l in lines[:5] if l.strip() and len(l.strip()) > 5]
        if len(first_lines) >= 2:
            combined = " ".join(first_lines[:2])
            if 20 < len(combined) < 350:
                return combined
        
        # Fallback
        for line in lines:
            line = line.strip()
            if 15 <= len(line) <= 300:
                return line
        
        return "Text Tender"


class FormTenderParser(BaseTenderParser):
    """Parse tender from structured form input."""

    def __init__(self, form_data: Dict[str, Any]):
        """
        Initialize form parser.
        
        Args:
            form_data: Dictionary containing tender form data with keys:
                - title: Tender title
                - description: Tender description/overview
                - requirements: Technical requirements
                - fleet_details: Fleet specifications
                - timeline: Project timeline
                - budget: Budget information
        """
        self.form_data = form_data

    def parse(self) -> TenderDocument:
        """
        Parse form-submitted tender.
        
        Returns:
            TenderDocument: Parsed tender
            
        Raises:
            TenderParsingError: If required fields are missing
        """
        # Validate required fields
        if 'title' not in self.form_data or not self.form_data['title'].strip():
            raise TenderParsingError("Tender title is required")
        
        # Build raw content from form fields
        raw_content_parts = []
        for key, value in self.form_data.items():
            if value:
                raw_content_parts.append(f"{key.upper()}:\n{value}\n")
        raw_content = "\n".join(raw_content_parts)
        
        # Parse form data into structured format
        technical_reqs = {}
        if 'requirements' in self.form_data:
            technical_reqs = self._parse_form_requirements(self.form_data['requirements'])
        
        fleet = {}
        if 'fleet_details' in self.form_data:
            fleet = self._parse_form_fleet(self.form_data['fleet_details'])
        
        logger.info(f"Form tender parsed. Title: {self.form_data['title']}")
        
        return TenderDocument(
            source_type=TenderSourceType.FORM,
            title=self.form_data['title'].strip(),
            raw_content=raw_content,
            sections={
                'overview': self.form_data.get('description', ''),
                'requirements': self.form_data.get('requirements', ''),
            },
            technical_requirements=technical_reqs,
            fleet_details=fleet,
            timeline=self.form_data.get('timeline'),
            budget_info=self.form_data.get('budget'),
            tender_no=self.form_data.get('tender_no'),
            address=self.form_data.get('address'),
            email=self.form_data.get('email'),
            phone_number=self.form_data.get('phone_number'),
            bid_validity=self.form_data.get('bid_validity'),
            metadata={'form_fields': list(self.form_data.keys())}
        )

    @staticmethod
    def _parse_form_requirements(requirements_text: str) -> Dict[str, Any]:
        """Parse technical requirements from form field."""
        # Split by common delimiters and structure
        return {
            'specifications': [r.strip() for r in requirements_text.split(',') if r.strip()],
            'compliance': [],
            'performance': []
        }

    @staticmethod
    def _parse_form_fleet(fleet_text: str) -> Dict[str, Any]:
        """Parse fleet details from form field."""
        return {
            'vehicles': [v.strip() for v in fleet_text.split(',') if v.strip()],
            'equipment': [],
            'specifications': []
        }


class LLMTenderParser:
    """
    LLM-based tender parser using Ollama for superior accuracy.
    Extracts structured information from any tender format using natural language understanding.
    """

    # Extraction prompt template - SIMPLIFIED FOR SPEED
    LLM_EXTRACTION_PROMPT = """Extract key information from this tender. Return ONLY valid JSON.

TENDER:
{tender_content}

Return JSON with: title, overview (brief), requirements (array), timeline, budget
Example: {{"title": "...", "overview": "...", "requirements": ["req1", "req2"], "timeline": "...", "budget": "..."}}"""

    @staticmethod
    def parse_with_llm(content: str, source_type: TenderSourceType) -> TenderDocument:
        """
        Parse tender using LLM for intelligent extraction.
        
        Args:
            content: Tender content (PDF text, plain text, or form data)
            source_type: Type of tender source
            
        Returns:
            TenderDocument: Parsed tender with LLM-extracted info
            
        Raises:
            TenderParsingError: If parsing fails
        """
        try:
            from app.ai_service.model_manager import get_model_manager, ModelConnectionError
            
            # Prepare content (truncate if too long)
            if len(content) > 6000:
                content = content[:6000] + "\n[... content truncated ...]"
            
            # Build the prompt
            prompt = LLMTenderParser.LLM_EXTRACTION_PROMPT.format(tender_content=content)
            
            # Get model manager and ensure a model is loaded
            model_manager = get_model_manager()
            if not model_manager.get_current_model():
                logger.info("Loading LLM for tender parsing...")
                model_manager.load_model()
            
            # Call the LLM with shorter timeout for Step 1 parsing
            logger.info(f"Parsing tender with LLM ({model_manager.get_current_model().name})...")
            response = model_manager.generate(
                prompt=prompt,
                temperature=0.1,  # Very low for consistent parsing
                max_tokens=800,  # Reduced for faster response
                timeout=45,  # 45 seconds max for tender parsing in Step 1
                system_prompt="Output only valid JSON."
            )
            
            # Parse the JSON response
            extracted = LLMTenderParser._parse_json_response(response)
            
            # Create TenderDocument from extracted data
            title = extracted.get('title', 'Untitled Tender').strip() or 'Untitled Tender'
            
            # Handle simplified structure
            requirements = extracted.get('requirements', [])
            if isinstance(requirements, list):
                tech_reqs = {'specifications': requirements, 'compliance': [], 'performance': []}
            else:
                tech_reqs = requirements if isinstance(requirements, dict) else {'specifications': [], 'compliance': [], 'performance': []}
            
            return TenderDocument(
                source_type=source_type,
                title=title,
                raw_content=content,
                sections={
                    'overview': extracted.get('overview', ''),
                    'requirements': str(extracted.get('requirements', ''))
                },
                technical_requirements=tech_reqs,
                fleet_details=extracted.get('fleet_details', {'vehicles': [], 'equipment': [], 'specifications': []}),
                timeline=extracted.get('timeline'),
                budget_info=extracted.get('budget'),
                metadata={
                    'source': source_type.value,
                    'extraction_method': 'llm',
                    'overview': extracted.get('overview', '')
                }
            )
        
        except (ModelConnectionError, ImportError) as e:
            logger.warning(f"LLM parsing unavailable: {str(e)}. Falling back to regex parser.")
            raise TenderParsingError(f"LLM parsing failed: {str(e)}")
        except json.JSONDecodeError as e:
            logger.warning(f"Failed to parse LLM response: {str(e)}. Falling back to regex parser.")
            raise TenderParsingError(f"Failed to parse LLM response: {str(e)}")
        except Exception as e:
            logger.error(f"LLM parsing error: {str(e)}")
            raise TenderParsingError(f"LLM parsing failed: {str(e)}")

    @staticmethod
    def _parse_json_response(response: str) -> Dict[str, Any]:
        """
        Parse JSON from LLM response.
        Handles markdown code blocks and extracts valid JSON.
        
        Args:
            response: LLM response text
            
        Returns:
            Dict: Parsed JSON
            
        Raises:
            json.JSONDecodeError: If JSON parsing fails
        """
        # Remove markdown code blocks
        if "```json" in response:
            start = response.find("```json") + 7
            end = response.find("```", start)
            response = response[start:end]
        elif "```" in response:
            start = response.find("```") + 3
            end = response.find("```", start)
            response = response[start:end]
        
        # Extract JSON object
        start = response.find('{')
        end = response.rfind('}') + 1
        if start >= 0 and end > start:
            response = response[start:end]
        
        return json.loads(response.strip())


class TenderParserFactory:
    """
    Factory class for creating appropriate tender parser based on input type.
    Handles parser creation, error handling, and logging.
    """

    @staticmethod
    def parse_pdf(file_bytes: bytes) -> TenderDocument:
        """
        Parse PDF tender using enhanced multi-method extraction:
        - pypdf for text extraction
        - pdfplumber for table extraction
        - PaddleOCR for scanned PDFs (fallback)
        
        Args:
            file_bytes: PDF file content as bytes
            
        Returns:
            TenderDocument: Parsed tender
            
        Raises:
            TenderParsingError: If parsing fails
        """
        try:
            logger.info("Parsing PDF with enhanced multi-method extraction...")
            parser = PDFTenderParser(file_bytes)
            return parser.parse()
        
        except TenderParsingError:
            raise
        except Exception as e:
            logger.error(f"PDF processing failed: {str(e)}")
            raise TenderParsingError(f"Failed to parse PDF: {str(e)}")

    @staticmethod
    def parse_text(text_content: str, title: Optional[str] = None) -> TenderDocument:
        """
        Parse text tender (email or plain text) using regex (fast parsing for Step 1).
        LLM parsing deferred to Step 3 (structure design).
        
        Args:
            text_content: Tender content as text
            title: Optional title to override extraction
            
        Returns:
            TenderDocument: Parsed tender
            
        Raises:
            TenderParsingError: If parsing fails
        """
        try:
            if not text_content or not text_content.strip():
                raise TenderParsingError("Text content cannot be empty")
            
            # Use regex-based parsing only for speed (Step 1)
            logger.info("Parsing text tender with regex (fast)...")
            parser = TextTenderParser(text_content, title)
            return parser.parse()
        
        except TenderParsingError:
            raise
        except Exception as e:
            logger.error(f"Text parsing failed: {str(e)}")
            raise TenderParsingError(f"Failed to parse text: {str(e)}")

    @staticmethod
    def parse_form(form_data: Dict[str, Any]) -> TenderDocument:
        """
        Parse form-submitted tender using regex (fast parsing for Step 1).
        LLM parsing deferred to Step 3 (structure design).
        
        Args:
            form_data: Dictionary with tender form data
            
        Returns:
            TenderDocument: Parsed tender
            
        Raises:
            TenderParsingError: If parsing fails
        """
        try:
            # Use regex-based parsing only for speed (Step 1)
            logger.info("Parsing form tender with regex (fast)...")
            parser = FormTenderParser(form_data)
            return parser.parse()
        
        except TenderParsingError:
            raise
        except Exception as e:
            logger.error(f"Form parsing failed: {str(e)}")
            raise TenderParsingError(f"Failed to parse form: {str(e)}")

    @staticmethod
    def parse_docx(file_bytes: bytes) -> TenderDocument:
        """
        Parse DOCX tender document.
        Extracts text from all paragraphs and tables in the DOCX file.
        
        Args:
            file_bytes: DOCX file content as bytes
            
        Returns:
            TenderDocument: Parsed tender
            
        Raises:
            TenderParsingError: If parsing fails
        """
        try:
            if not PYTHON_DOCX_AVAILABLE:
                raise TenderParsingError("python-docx is not installed. Please install it using: pip install python-docx")
            
            logger.info("Parsing DOCX document...")
            
            # Load the DOCX file from bytes
            docx_file = io.BytesIO(file_bytes)
            doc = Document(docx_file)
            
            # Extract all text from paragraphs
            text_content = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text)
                    if row_text:
                        text_content.append(' | '.join(row_text))
            
            # Combine all text
            full_text = '\n'.join(text_content)
            
            if not full_text.strip():
                raise TenderParsingError("No text content found in DOCX file")
            
            # Use text parser on extracted content
            return TenderParserFactory.parse_text(full_text, title=None)
        
        except TenderParsingError:
            raise
        except Exception as e:
            logger.error(f"DOCX parsing failed: {str(e)}")
            raise TenderParsingError(f"Failed to parse DOCX: {str(e)}")

    @staticmethod
    def parse_txt(file_bytes: bytes) -> TenderDocument:
        """
        Parse TXT tender document.
        Decodes the text file and processes it like text input.
        
        Args:
            file_bytes: TXT file content as bytes
            
        Returns:
            TenderDocument: Parsed tender
            
        Raises:
            TenderParsingError: If parsing fails
        """
        try:
            logger.info("Parsing TXT document...")
            
            # Try multiple encodings
            text_content = None
            for encoding in ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']:
                try:
                    text_content = file_bytes.decode(encoding)
                    break
                except UnicodeDecodeError:
                    continue
            
            if text_content is None:
                raise TenderParsingError("Could not decode TXT file with any supported encoding")
            
            if not text_content.strip():
                raise TenderParsingError("TXT file is empty")
            
            # Use text parser on the content
            return TenderParserFactory.parse_text(text_content, title=None)
        
        except TenderParsingError:
            raise
        except Exception as e:
            logger.error(f"TXT parsing failed: {str(e)}")
            raise TenderParsingError(f"Failed to parse TXT: {str(e)}")

    @staticmethod
    def parse_file(file_bytes: bytes, filename: str) -> TenderDocument:
        """
        Parse a file based on its extension.
        Automatically detects the file type and uses the appropriate parser.
        
        Args:
            file_bytes: File content as bytes
            filename: Original filename (used to detect file type)
            
        Returns:
            TenderDocument: Parsed tender
            
        Raises:
            TenderParsingError: If file type is not supported
        """
        # Get file extension
        file_ext = filename.split('.')[-1].lower() if '.' in filename else ''
        
        logger.info(f"Detecting file type from extension: .{file_ext}")
        
        if file_ext == 'pdf':
            return TenderParserFactory.parse_pdf(file_bytes)
        elif file_ext == 'docx' or file_ext == 'doc':
            return TenderParserFactory.parse_docx(file_bytes)
        elif file_ext == 'txt':
            return TenderParserFactory.parse_txt(file_bytes)
        else:
            raise TenderParsingError(f"Unsupported file type: .{file_ext}. Supported types: PDF, DOCX, TXT")

    @staticmethod
    def _form_to_text(form_data: Dict[str, Any]) -> str:
        """
        Convert form data to text for LLM parsing.
        
        Args:
            form_data: Form fields dictionary
            
        Returns:
            str: Formatted text content
        """
        parts = []
        
        if form_data.get('title'):
            parts.append(f"TENDER TITLE: {form_data['title']}")
        
        if form_data.get('description'):
            parts.append(f"DESCRIPTION:\n{form_data['description']}")
        
        if form_data.get('requirements'):
            parts.append(f"REQUIREMENTS:\n{form_data['requirements']}")
        
        if form_data.get('fleet_details'):
            parts.append(f"FLEET DETAILS:\n{form_data['fleet_details']}")
        
        if form_data.get('timeline'):
            parts.append(f"TIMELINE: {form_data['timeline']}")
        
        if form_data.get('budget'):
            parts.append(f"BUDGET: {form_data['budget']}")
        
        return "\n\n".join(parts)
