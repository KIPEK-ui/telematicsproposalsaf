"""
Form Detector Service
Detects form fields, tables, and text fields in tender documents.
Supports DOCX and PDF formats.
✅ ENHANCED: Caches form detection results to avoid re-scanning
"""

import logging
import hashlib
import json
from typing import Dict, List, Any, Optional, Tuple
from dataclasses import dataclass, field, asdict
from pathlib import Path
from enum import Enum
from datetime import datetime

logger = logging.getLogger(__name__)

try:
    from docx import Document
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import pdfplumber
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False


class FieldType(str, Enum):
    """Types of form fields."""
    TEXT = "text"
    TABLE = "table"
    TABLE_ROW = "table_row"
    PARAGRAPH = "paragraph"
    MULTILINE = "multiline"
    CHECKBOX = "checkbox"
    LIST = "list"
    UNKNOWN = "unknown"


@dataclass
class FormField:
    """Represents a detected form field."""
    field_id: str
    field_type: FieldType
    label: str
    placeholder: str = ""
    location: Dict[str, Any] = field(default_factory=dict)  # paragraph index, table index, etc.
    estimated_length: int = 0  # chars
    parent_section: str = ""  # heading it belongs under
    table_context: Optional[Dict[str, Any]] = None  # if in a table
    confidence: float = 1.0  # 0-1 confidence in detection

    def to_dict(self) -> Dict[str, Any]:
        """Convert to dictionary."""
        return {
            'field_id': self.field_id,
            'field_type': self.field_type.value,
            'label': self.label,
            'placeholder': self.placeholder,
            'location': self.location,
            'estimated_length': self.estimated_length,
            'parent_section': self.parent_section,
            'table_context': self.table_context,
            'confidence': self.confidence
        }


@dataclass
class FormStructure:
    """Complete structure of detected forms in a document."""
    document_path: str
    document_type: str  # 'docx' or 'pdf'
    fields: List[FormField] = field(default_factory=list)
    tables: List[Dict[str, Any]] = field(default_factory=list)  # table structures
    sections: List[str] = field(default_factory=list)  # document headings/sections
    detected_at: str = ""  # timestamp

    def to_dict(self) -> Dict[str, Any]:
        """Convert to dictionary."""
        return {
            'document_path': self.document_path,
            'document_type': self.document_type,
            'fields': [f.to_dict() for f in self.fields],
            'tables': self.tables,
            'sections': self.sections,
            'detected_at': self.detected_at
        }


class FormDetector:
    """
    Detects form fields, tables, and fillable areas in tender documents.
    Works with DOCX and PDF formats.
    ✅ ENHANCED: Caches form detection results to avoid re-scanning same files
    """

    def __init__(self):
        """Initialize form detector."""
        if not DOCX_AVAILABLE:
            logger.warning("python-docx not available. DOCX support disabled.")
        if not PDF_AVAILABLE:
            logger.warning("pdfplumber not available. PDF support disabled.")
        
        # ✅ NEW: Form detection cache
        self.cache_dir = Path("data/.form_detection_cache")
        self.cache_dir.mkdir(parents=True, exist_ok=True)
        self.cache_file = self.cache_dir / "form_detection_cache.json"
        self._form_cache: Dict[str, Dict[str, Any]] = {}
        self._load_form_cache()
        
        logger.info("FormDetector initialized")

    def _get_file_hash(self, file_path: Path) -> str:
        """Generate hash of file for cache validation."""
        try:
            with open(file_path, 'rb') as f:
                return hashlib.md5(f.read()).hexdigest()
        except Exception:
            return ""

    def _load_form_cache(self) -> None:
        """Load form detection cache from disk."""
        try:
            if self.cache_file.exists():
                with open(self.cache_file, 'r', encoding='utf-8') as f:
                    self._form_cache = json.load(f)
                logger.debug(f"Loaded form detection cache with {len(self._form_cache)} entries")
        except Exception as e:
            logger.warning(f"Failed to load form detection cache: {e}")
            self._form_cache = {}

    def _save_form_cache(self) -> None:
        """Save form detection cache to disk."""
        try:
            with open(self.cache_file, 'w', encoding='utf-8') as f:
                json.dump(self._form_cache, f, indent=2, default=str)
            logger.debug("Saved form detection cache")
        except Exception as e:
            logger.warning(f"Failed to save form detection cache: {e}")

    def _is_form_cached(self, file_path: Path) -> Optional[FormStructure]:
        """Check if form structure was already detected and is still valid."""
        try:
            file_hash = self._get_file_hash(file_path)
            cache_key = str(file_path.absolute())
            
            if cache_key in self._form_cache:
                cached = self._form_cache[cache_key]
                if cached.get('file_hash') == file_hash and cached.get('form_structure'):
                    logger.debug(f"Form cache hit for {file_path.name}")
                    # Reconstruct FormStructure from cached data
                    cached_data = cached['form_structure']
                    return FormStructure(
                        document_path=cached_data['document_path'],
                        document_type=cached_data['document_type'],
                        fields=cached_data['fields'],
                        tables=cached_data['tables'],
                        sections=cached_data['sections']
                    )
        except Exception as e:
            logger.debug(f"Form cache check failed: {e}")
        
        return None

    def _cache_form_structure(self, file_path: Path, form_structure: FormStructure) -> None:
        """Cache form detection result for future use."""
        try:
            file_hash = self._get_file_hash(file_path)
            cache_key = str(file_path.absolute())
            
            self._form_cache[cache_key] = {
                'file_hash': file_hash,
                'form_structure': asdict(form_structure),
                'cached_at': datetime.now().isoformat()
            }
            self._save_form_cache()
        except Exception as e:
            logger.warning(f"Failed to cache form structure: {e}")

    def detect_form_structure(self, file_path: str) -> FormStructure:
        """
        Detect form structure in a tender document.
        ✅ ENHANCED: Checks cache first to avoid re-scanning same file

        Args:
            file_path: Path to tender document (DOCX or PDF)

        Returns:
            FormStructure: Detected fields, tables, and sections
        """
        file_path = Path(file_path)

        if not file_path.exists():
            logger.error(f"File not found: {file_path}")
            return FormStructure(
                document_path=str(file_path),
                document_type="unknown",
                fields=[],
                tables=[],
                sections=[]
            )

        # ✅ NEW: Check cache first
        cached_form = self._is_form_cached(file_path)
        if cached_form is not None:
            return cached_form

        suffix = file_path.suffix.lower()

        if suffix == ".docx":
            form_structure = self._detect_docx_form(file_path)
        elif suffix == ".pdf":
            form_structure = self._detect_pdf_form(file_path)
        else:
            logger.warning(f"Unsupported file type: {suffix}")
            form_structure = FormStructure(
                document_path=str(file_path),
                document_type="unknown",
                fields=[],
                tables=[],
                sections=[]
            )
        
        # ✅ NEW: Cache the result
        self._cache_form_structure(file_path, form_structure)
        return form_structure

    def _detect_docx_form(self, file_path: Path) -> FormStructure:
        """Detect form structure in DOCX document."""
        if not DOCX_AVAILABLE:
            return FormStructure(
                document_path=str(file_path),
                document_type="docx",
                fields=[],
                tables=[],
                sections=[]
            )

        try:
            doc = Document(file_path)
            form_structure = FormStructure(
                document_path=str(file_path),
                document_type="docx"
            )

            # Extract sections (headings)
            form_structure.sections = self._extract_sections(doc)

            # Extract form fields and tables
            field_id_counter = 0
            current_section = ""

            for para_idx, para in enumerate(doc.paragraphs):
                # Update current section if this is a heading
                if para.style.name.startswith("Heading"):
                    current_section = para.text.strip()

                # Detect form fields in paragraph
                field = self._detect_field_in_paragraph(
                    para, para_idx, current_section, field_id_counter
                )
                if field:
                    form_structure.fields.append(field)
                    field_id_counter += 1

            # Extract tables
            for table_idx, table in enumerate(doc.tables):
                table_info = self._analyze_table(table, table_idx, current_section)
                form_structure.tables.append(table_info)

                # Add table fields
                for row_idx, row in enumerate(table.rows):
                    for cell_idx, cell in enumerate(row.cells):
                        text = cell.text.strip()
                        if text and self._is_fillable_cell(text):
                            field = FormField(
                                field_id=f"table_{table_idx}_row_{row_idx}_cell_{cell_idx}",
                                field_type=FieldType.TABLE_ROW,
                                label=text,
                                location={
                                    "table_index": table_idx,
                                    "row_index": row_idx,
                                    "cell_index": cell_idx
                                },
                                parent_section=current_section,
                                table_context={
                                    "table_index": table_idx,
                                    "rows": len(table.rows),
                                    "cols": len(table.columns)
                                }
                            )
                            form_structure.fields.append(field)
                            field_id_counter += 1

            logger.info(f"Detected {len(form_structure.fields)} fields in DOCX: {file_path}")
            return form_structure

        except Exception as e:
            logger.error(f"Error detecting DOCX form: {e}")
            return FormStructure(
                document_path=str(file_path),
                document_type="docx",
                fields=[],
                tables=[],
                sections=[]
            )

    def _detect_pdf_form(self, file_path: Path) -> FormStructure:
        """Detect form structure in PDF document."""
        if not PDF_AVAILABLE:
            return FormStructure(
                document_path=str(file_path),
                document_type="pdf",
                fields=[],
                tables=[],
                sections=[]
            )

        try:
            form_structure = FormStructure(
                document_path=str(file_path),
                document_type="pdf"
            )

            with pdfplumber.open(file_path) as pdf:
                # Detect form fields from PDF metadata
                if pdf.metadata and "AcroForm" in str(pdf.metadata):
                    logger.info("PDF contains form fields")

                # Extract text and detect fillable areas
                field_id_counter = 0
                current_section = ""

                for page_num, page in enumerate(pdf.pages):
                    text = page.extract_text()
                    if text:
                        # Simple heuristic: look for lines with colons (label: _____)
                        lines = text.split("\n")
                        for line_idx, line in enumerate(lines):
                            if ":" in line and len(line) > 5:
                                parts = line.split(":")
                                if len(parts) == 2:
                                    label = parts[0].strip()
                                    value_part = parts[1].strip()

                                    # Detect if this looks like a fillable field
                                    if self._is_fillable_text(label):
                                        field = FormField(
                                            field_id=f"pdf_p{page_num}_f{field_id_counter}",
                                            field_type=FieldType.TEXT,
                                            label=label,
                                            placeholder=value_part,
                                            location={
                                                "page": page_num,
                                                "line_index": line_idx
                                            },
                                            parent_section=current_section,
                                            estimated_length=100,
                                            confidence=0.7
                                        )
                                        form_structure.fields.append(field)
                                        field_id_counter += 1

            logger.info(f"Detected {len(form_structure.fields)} fields in PDF: {file_path}")
            return form_structure

        except Exception as e:
            logger.error(f"Error detecting PDF form: {e}")
            return FormStructure(
                document_path=str(file_path),
                document_type="pdf",
                fields=[],
                tables=[],
                sections=[]
            )

    def _extract_sections(self, doc: Document) -> List[str]:
        """Extract section headings from document."""
        sections = []
        for para in doc.paragraphs:
            if para.style.name.startswith("Heading"):
                text = para.text.strip()
                if text:
                    sections.append(text)
        return sections

    def _detect_field_in_paragraph(
        self, para, para_idx: int, current_section: str, field_counter: int
    ) -> Optional[FormField]:
        """Detect form field in a paragraph."""
        text = para.text.strip()

        # Look for field patterns like "Field Name: _____" or "[Field Name]"
        if not text or len(text) < 3:
            return None

        # Pattern 1: "Label: _____"
        if ":" in text:
            parts = text.split(":", 1)
            label = parts[0].strip()
            value_part = parts[1].strip() if len(parts) > 1 else ""

            if self._is_fillable_text(label) and (
                "_" in value_part or value_part == "" or len(value_part) < 50
            ):
                return FormField(
                    field_id=f"docx_p{para_idx}_f{field_counter}",
                    field_type=FieldType.TEXT,
                    label=label,
                    placeholder=value_part,
                    location={"paragraph_index": para_idx},
                    parent_section=current_section,
                    estimated_length=200
                )

        # Pattern 2: "[Field Name]" or "{Field Name}"
        if ("[" in text and "]" in text) or ("{" in text and "}" in text):
            # Extract field name
            if "[" in text:
                field_name = text.split("[")[1].split("]")[0]
            else:
                field_name = text.split("{")[1].split("}")[0]

            if field_name.strip():
                return FormField(
                    field_id=f"docx_p{para_idx}_f{field_counter}",
                    field_type=FieldType.PARAGRAPH,
                    label=field_name.strip(),
                    location={"paragraph_index": para_idx},
                    parent_section=current_section,
                    estimated_length=300
                )

        return None

    def _analyze_table(
        self, table, table_idx: int, current_section: str
    ) -> Dict[str, Any]:
        """Analyze table structure."""
        return {
            "table_index": table_idx,
            "rows": len(table.rows),
            "columns": len(table.columns),
            "parent_section": current_section,
            "description": f"Table with {len(table.rows)} rows and {len(table.columns)} columns"
        }

    def _is_fillable_cell(self, text: str) -> bool:
        """Check if a cell looks like it should be fillable."""
        # Empty cells or cells with only underscores
        if not text or text.strip() == "":
            return False
        if "_" * 3 in text:
            return True
        # Short cells that aren't typical content
        if len(text) < 50 and not any(
            phrase in text.lower()
            for phrase in ["the", "and", "this", "that", "with", "for"]
        ):
            return True
        return False

    def _is_fillable_text(self, text: str) -> bool:
        """Check if text looks like a form field label."""
        if not text or len(text) < 2:
            return False
        # Avoid section headings
        if any(text.lower().endswith(suffix) for suffix in [" summary", " details", " section", " information"]):
            return True
        # Check for typical field label patterns
        if any(pattern in text.lower() for pattern in [
            "name", "date", "address", "phone", "email", "company",
            "title", "description", "amount", "signature", "approved",
            "cost", "price", "budget", "timeline", "schedule", "start",
            "end", "duration", "specification", "requirement", "notes",
            "comments", "reference", "number", "code", "id"
        ]):
            return True
        return False


class FormDetectionError(Exception):
    """Raised when form detection fails."""
    pass
