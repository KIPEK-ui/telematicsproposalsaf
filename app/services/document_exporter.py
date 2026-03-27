"""
Document Exporter Service
Converts proposal content to professionally formatted .docx files.
Includes Safaricom branding, styling, and structure.
Also supports filling tender templates with proposal content.
"""

import logging
from typing import Dict, Any, Optional, Tuple
from datetime import datetime
from io import BytesIO
from pathlib import Path
from zipfile import ZipFile

logger = logging.getLogger(__name__)

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class DocumentExportError(Exception):
    """Raised when document export fails."""
    pass


class DocumentExporter:
    """
    Exports proposal content to professionally formatted .docx documents.
    Features:
    - Safaricom branding (logo, colors)
    - Professional styling
    - Structured sections with TOC
    - Multi-page support
    """

    # Safaricom brand colors (stored as RGB tuples for flexibility)
    BRAND_COLORS = {
        'primary': (255, 99, 0),      # Safaricom orange
        'secondary': (51, 51, 51),     # Dark gray
        'accent': (200, 200, 200)      # Light gray
    }

    SECTION_HEADINGS = {
        'executive_summary': ('Executive Summary', 1),
        'security': ('Data Security & Compliance', 1),
        'technical_approach': ('Technical Approach', 1),
        'implementation_timeline': ('Implementation Timeline', 1),
        'pricing': ('Pricing & Commercial Terms', 1),
        'terms_conditions': ('Terms & Conditions', 1),
        'fleet_details': ('Fleet & Equipment Details', 1),
        'compliance_assurance': ('Compliance & Quality Assurance', 1),
        'financial_proposal': ('Financial Proposal', 1),
        'project_management': ('Project Management', 1),
    }
    
    # Preferred section order for logical flow
    SECTION_ORDER = [
        'executive_summary',
        'security',
        'technical_approach',
        'implementation_timeline',
        'pricing',
        'terms_conditions',
        'compliance_assurance',
        'fleet_details',
        'financial_proposal',
        'project_management',
    ]

    def __init__(self):
        """Initialize document exporter."""
        if not DOCX_AVAILABLE:
            raise DocumentExportError(
                "python-docx not available. Install: pip install python-docx"
            )
        # Set branding paths
        branding_dir = Path(__file__).parent.parent.parent / "data" / "branding"
        self.logo_path = branding_dir / "LOGO.jpg"
        self.details_path = branding_dir / "DETAILS.jpg"
        logger.info("DocumentExporter initialized with branding assets")

    def export_to_docx(
        self,
        proposal_content: Dict[str, str],
        org_data: Dict[str, Any],
        tender_title: str = "Proposal Document",
        tender_reference: str = ""
    ) -> bytes:
        """
        Export proposal to .docx format.
        
        Args:
            proposal_content: Dictionary with proposal sections
            org_data: Organization information (name, industry, contact)
            tender_title: Title of the tender/proposal
            tender_reference: Tender reference number (e.g., KRA/HQS/DP-016/2025-2026)
        
        Returns:
            bytes: .docx file content
            
        Raises:
            DocumentExportError: If export fails
        """
        try:
            doc = Document()
            
            # Store org_data on document object for use in _add_sections
            doc._org_data = org_data
            
            # Set document margins
            sections = doc.sections
            for section in sections:
                section.top_margin = Cm(2)
                section.bottom_margin = Cm(2)
                section.left_margin = Cm(2)
                section.right_margin = Cm(2)
            
            # Add header/title section (cover page only, no org info here)
            self._add_header(doc, org_data, tender_title, tender_reference)
            
            # Add proposal_content sections (org info box + all content sections)
            self._add_sections(doc, proposal_content, org_data)
            
            # Add footer
            self._add_footer(doc, org_data)
            
            # Convert to bytes
            output = BytesIO()
            doc.save(output)
            output.seek(0)
            
            logger.info("Document exported successfully to .docx")
            return output.getvalue()
        
        except Exception as e:
            logger.error(f"Document export failed: {str(e)}")
            raise DocumentExportError(f"Failed to export document: {str(e)}")

    def export_to_file(
        self,
        proposal_content: Dict[str, str],
        org_data: Dict[str, Any],
        output_path: str,
        tender_title: str = "Proposal Document",
        tender_reference: str = ""
    ) -> str:
        """
        Export proposal to .docx file on disk.
        
        Args:
            proposal_content: Dictionary with proposal sections
            org_data: Organization information
            output_path: Path to save .docx file
            tender_title: Title of the proposal
            tender_reference: Tender reference number
        
        Returns:
            str: Path to saved file
            
        Raises:
            DocumentExportError: If export fails
        """
        try:
            output_path = Path(output_path)
            output_path.parent.mkdir(parents=True, exist_ok=True)
            
            docx_bytes = self.export_to_docx(
                proposal_content, org_data, tender_title, tender_reference
            )
            
            with open(output_path, 'wb') as f:
                f.write(docx_bytes)
            
            logger.info(f"Document saved to {output_path}")
            return str(output_path)
        
        except Exception as e:
            logger.error(f"File export failed: {str(e)}")
            raise DocumentExportError(f"Failed to save file: {str(e)}")

    def export_as_filled_tender(
        self,
        original_tender_path: str,
        proposal_content: Dict[str, str],
        org_data: Dict[str, Any]
    ) -> bytes:
        """
        Export proposal as filled tender template (preserving original structure).
        
        Args:
            original_tender_path: Path to original tender document
            proposal_content: Dictionary with proposal sections
            org_data: Organization information
        
        Returns:
            bytes: Filled tender document content
            
        Raises:
            DocumentExportError: If filling fails
        """
        try:
            from .form_detector import FormDetector
            from .form_filler import FormFiller

            # Check if the tender file exists
            tender_path = Path(original_tender_path)
            if not tender_path.exists():
                logger.warning(f"Tender file not found: {original_tender_path}")
                raise DocumentExportError(f"Tender file not found: {original_tender_path}")

            logger.info(f"Attempting to fill tender from: {original_tender_path}")
            
            # Detect form structure in original tender
            detector = FormDetector()
            form_structure = detector.detect_form_structure(str(tender_path))
            logger.info(f"Detected form structure: {len(form_structure) if form_structure else 0} fields")

            # Fill the form with proposal content
            filler = FormFiller()
            filled_docx = filler.fill_form(
                str(tender_path),
                form_structure,
                proposal_content,
                org_data
            )
            
            if not filled_docx:
                raise DocumentExportError("Form filler returned empty document")

            logger.info("Filled tender exported successfully")
            return filled_docx

        except DocumentExportError:
            raise
        except Exception as e:
            logger.error(f"Tender filling failed: {str(e)}")
            raise DocumentExportError(f"Failed to fill tender: {str(e)}")

    def export_dual(
        self,
        proposal_content: Dict[str, str],
        org_data: Dict[str, Any],
        original_tender_path: Optional[str] = None,
        tender_title: str = "Proposal Document",
        tender_reference: str = ""
    ) -> Tuple[bytes, Optional[bytes]]:
        """
        Export both branded proposal and filled tender (if original provided).
        
        Args:
            proposal_content: Dictionary with proposal sections
            org_data: Organization information
            original_tender_path: Path to original tender (optional)
            tender_title: Title of the proposal
            tender_reference: Tender reference number
        
        Returns:
            Tuple[bytes, Optional[bytes]]: (branded_docx, filled_tender_docx or None)
        """
        try:
            # Always generate branded proposal
            branded = self.export_to_docx(
                proposal_content, org_data, tender_title, tender_reference
            )

            # Optionally fill original tender if provided
            filled_tender = None
            if original_tender_path and Path(original_tender_path).exists():
                try:
                    filled_tender = self.export_as_filled_tender(
                        original_tender_path,
                        proposal_content,
                        org_data
                    )
                    logger.info("Both branded and filled tender generated")
                except Exception as e:
                    logger.warning(f"Could not generate filled tender: {e}. Returning branded only.")

            return (branded, filled_tender)

        except Exception as e:
            logger.error(f"Dual export failed: {str(e)}")
            raise DocumentExportError(f"Failed to export documents: {str(e)}")

    def export_dual_as_zip(
        self,
        proposal_content: Dict[str, str],
        org_data: Dict[str, Any],
        original_tender_path: Optional[str] = None,
        tender_title: str = "Proposal Document",
        tender_reference: str = ""
    ) -> bytes:
        """
        Export both branded and filled tender as a ZIP archive.
        
        Args:
            proposal_content: Dictionary with proposal sections
            org_data: Organization information
            original_tender_path: Path to original tender (optional)
            tender_title: Title of the proposal
            tender_reference: Tender reference number
        
        Returns:
            bytes: ZIP file content containing both documents
        """
        try:
            branded, filled_tender = self.export_dual(
                proposal_content,
                org_data,
                original_tender_path,
                tender_title,
                tender_reference
            )

            # Create ZIP archive
            zip_buffer = BytesIO()
            with ZipFile(zip_buffer, 'w') as zip_file:
                # Add branded proposal
                org_safe = "".join(c if c.isalnum() or c == " " else "" for c in org_data.get('name', 'Org'))[:20].replace(" ", "_")
                tender_safe = "".join(c if c.isalnum() or c == " " else "" for c in tender_title)[:20].replace(" ", "_")
                
                branded_filename = f"Proposal_Branded_{org_safe}_{tender_safe}.docx"
                zip_file.writestr(branded_filename, branded)

                # Add filled tender if available
                if filled_tender:
                    filled_filename = f"Proposal_Filled_Tender_{org_safe}_{tender_safe}.docx"
                    zip_file.writestr(filled_filename, filled_tender)

            zip_buffer.seek(0)
            logger.info("Dual export as ZIP completed")
            return zip_buffer.getvalue()

        except Exception as e:
            logger.error(f"ZIP export failed: {str(e)}")
            raise DocumentExportError(f"Failed to create ZIP export: {str(e)}")

    # ========================
    # Document Building Methods
    # ========================

    def _add_header(
        self,
        doc: Document,
        org_data: Dict[str, Any],
        tender_title: str,
        tender_reference: str = ""
    ) -> None:
        """Add title and header section with Safaricom branding - cover page only, generated once."""
        # Create clean cover page with proper spacing
        
        # Top spacing (push content down)
        for _ in range(3):
            doc.add_paragraph()
        
        # Add logo if available
        if self.logo_path.exists():
            try:
                logo_para = doc.add_paragraph()
                logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                logo_run = logo_para.add_run()
                logo_run.add_picture(str(self.logo_path), width=Inches(2.5))
            except Exception as e:
                logger.warning(f"Could not add logo: {e}")
        
        # Add spacing after logo
        for _ in range(2):
            doc.add_paragraph()
        
        # Title
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.add_run("PROPOSAL")
        title_run.font.size = Pt(32)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(*self.BRAND_COLORS['primary'])
        
        # Spacing
        doc.add_paragraph()
        
        # Tender title
        subtitle = doc.add_paragraph(tender_title)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if subtitle.runs:
            subtitle_run = subtitle.runs[0]
            subtitle_run.font.size = Pt(18)
            subtitle_run.font.bold = True
            subtitle_run.font.color.rgb = RGBColor(*self.BRAND_COLORS['secondary'])
        
        # Spacing
        for _ in range(2):
            doc.add_paragraph()
        
        # Tender reference (if provided)
        if tender_reference and tender_reference.strip():
            ref_para = doc.add_paragraph(f"Tender Reference: {tender_reference}")
            ref_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if ref_para.runs:
                ref_run = ref_para.runs[0]
                ref_run.font.size = Pt(11)
                ref_run.font.bold = True
                ref_run.font.color.rgb = RGBColor(*self.BRAND_COLORS['primary'])
            doc.add_paragraph()
        
        # Organization name
        org_para = doc.add_paragraph(org_data.get('name', 'Organization'))
        org_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if org_para.runs:
            org_run = org_para.runs[0]
            org_run.font.size = Pt(14)
            org_run.font.bold = True
        
        # Date
        date_para = doc.add_paragraph(f"Prepared: {datetime.now().strftime('%B %d, %Y')}")
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if date_para.runs:
            date_run = date_para.runs[0]
            date_run.font.size = Pt(11)
            date_run.font.italic = True
            date_run.font.color.rgb = RGBColor(100, 100, 100)
        
        # Bottom spacing
        for _ in range(3):
            doc.add_paragraph()
        
        # Add page break after cover page - content follows on next page
        doc.add_page_break()

    def _add_sections(
        self,
        doc: Document,
        proposal_content: Dict[str, str],
        org_data: Dict[str, Any]
    ) -> None:
        """Add organization info box, then proposal sections in logical order with styled headers and plain content."""
        # Keys to skip - these are metadata, not proposal sections
        METADATA_KEYS = {
            'Cover Page', 'cover_page',
            'Section Order', 'section_order',
            'Design Rationale', 'design_rationale',
            'Success Factors', 'success_factors',
            'Proposal Structure', 'proposal_structure',
            'Section Order', 'section_names'
        }
        
        # Create organization info box at the start (before any content sections)
        title = doc.add_heading("Organization Information", level=2)
        if title.runs:
            title.runs[0].font.color.rgb = RGBColor(*self.BRAND_COLORS['primary'])
            title.runs[0].font.size = Pt(18)
        
        # Create table with proper structure
        table = doc.add_table(rows=3, cols=2)
        table.style = 'Light Grid Accent 1'
        
        # Set column widths
        for row in table.rows:
            row.cells[0].width = Inches(2)
            row.cells[1].width = Inches(3.5)
        
        # org_data is now passed as parameter
        
        # Row 1: Name
        self._set_cell_text(table.rows[0].cells[0], "Organization Name:", bold=True)
        self._set_cell_text(table.rows[0].cells[1], org_data.get('name', 'N/A'))
        
        # Row 2: Industry
        self._set_cell_text(table.rows[1].cells[0], "Industry:", bold=True)
        self._set_cell_text(table.rows[1].cells[1], org_data.get('industry', 'N/A'))
        
        # Row 3: Contact Info combined
        self._set_cell_text(table.rows[2].cells[0], "Contact:", bold=True)
        contact_text = f"Email: {org_data.get('contact_email', 'N/A')}\nPhone: {org_data.get('contact_phone', 'N/A')}"
        self._set_cell_text(table.rows[2].cells[1], contact_text)
        
        # Add spacing after organization info
        doc.add_paragraph()
        doc.add_page_break()
        
        # Filter and order sections - skip metadata keys
        ordered_content = []
        
        # First, add sections in preferred order
        for section_key in self.SECTION_ORDER:
            if section_key in proposal_content and section_key not in METADATA_KEYS:
                ordered_content.append((section_key, proposal_content[section_key]))
        
        # Then, add any remaining sections not in SECTION_ORDER (except metadata keys)
        for key, content in proposal_content.items():
            if key not in METADATA_KEYS and not any(k == key for k, _ in ordered_content):
                ordered_content.append((key, content))
        
        # Add each section with styled header and plain content
        for idx, (key, content) in enumerate(ordered_content):
            # Handle both string and dict content
            if isinstance(content, dict):
                # Convert dict to formatted string
                content_str = self._dict_to_text(content)
            else:
                # Already a string
                content_str = str(content) if content else ""
            
            # Only add if there's actual content
            if content_str and content_str.strip():
                # Add page break between sections (except after org info)
                if idx > 0:
                    doc.add_page_break()
                
                # Generate heading from key
                if key in self.SECTION_HEADINGS:
                    heading, level = self.SECTION_HEADINGS[key]
                else:
                    # Convert key to title case for dynamic sections
                    heading = key.replace('_', ' ').title()
                    level = 1
                
                # Add heading ONLY - styled with Safaricom branding
                heading_para = doc.add_heading(heading, level=level)
                heading_para.paragraph_format.space_before = Pt(0)
                heading_para.paragraph_format.space_after = Pt(12)
                
                # Format heading text - ORANGE, BOLD, LARGE
                if heading_para.runs:
                    for run in heading_para.runs:
                        run.font.color.rgb = RGBColor(*self.BRAND_COLORS['primary'])
                        run.font.size = Pt(18)
                        run.font.bold = True
                
                # Add decorative line under heading
                line_para = doc.add_paragraph()
                line_para.paragraph_format.space_before = Pt(0)
                line_para.paragraph_format.space_after = Pt(12)
                pPr = line_para._element.get_or_add_pPr()
                pBdr = OxmlElement('w:pBdr')
                bottom = OxmlElement('w:bottom')
                bottom.set(qn('w:val'), 'single')
                bottom.set(qn('w:sz'), '12')  # 1.5pt
                bottom.set(qn('w:space'), '1')
                bottom.set(qn('w:color'), 'FF6300')  # Safaricom orange
                pBdr.append(bottom)
                pPr.append(pBdr)
                
                # Add content with PLAIN styling (no color/size changes)
                paragraphs = content_str.split('\n\n')
                for para_text in paragraphs:
                    if para_text.strip():
                        para = doc.add_paragraph(para_text.strip())
                        para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
                        para.paragraph_format.space_before = Pt(6)
                        para.paragraph_format.space_after = Pt(6)
                        
                        # Plain text - default formatting, no color override
                        for run in para.runs:
                            run.font.size = Pt(11)
                            # Don't override color - keep default black

    def _add_footer(
        self,
        doc: Document,
        org_data: Dict[str, Any]
    ) -> None:
        """Add footer with Safaricom branding details and contact information."""
        # Add footer with details image if available
        section = doc.sections[0]
        footer = section.footer
        
        # Add details image if available
        if self.details_path.exists():
            try:
                # Clear default footer paragraphs
                for para in footer.paragraphs:
                    p = para._element
                    p.getparent().remove(p)
                
                # Add new paragraph with image
                footer_para = footer.add_paragraph()
                footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                footer_run = footer_para.add_run()
                footer_run.add_picture(str(self.details_path), width=Inches(6))
                logger.info("Added Safaricom details image to footer")
            except Exception as e:
                logger.warning(f"Could not add footer image: {e}")
                # Fallback to text footer
                footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
                footer_para.text = f"{org_data.get('name', 'Safaricom')} | Confidential | {datetime.now().strftime('%Y-%m-%d')}"
                footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in footer_para.runs:
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(128, 128, 128)
        else:
            # Fallback text footer if image not available
            footer_para = footer.paragraphs[0]
            footer_para.text = f"{org_data.get('name', 'Safaricom')} | Confidential | {datetime.now().strftime('%Y-%m-%d')}"
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in footer_para.runs:
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(128, 128, 128)

    @staticmethod
    def _shade_cell(cell, color):
        """Shade a table cell with background color."""
        shading_elm = OxmlElement('w:shd')
        # Convert RGB tuple to hex string (e.g., (255, 99, 0) -> 'FF6300')
        if isinstance(color, tuple) and len(color) == 3:
            r, g, b = color
            hex_color = f'{r:02X}{g:02X}{b:02X}'
        else:
            hex_color = 'CCCCCC'  # Fallback to light gray
        
        shading_elm.set(qn('w:fill'), hex_color)
        cell._element.get_or_add_tcPr().append(shading_elm)
    
    @staticmethod
    def _set_cell_text(cell, text: str, bold: bool = False):
        """Set text in a table cell with optional formatting."""
        cell.text = text
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                if bold:
                    run.font.bold = True
                run.font.size = Pt(11)

    @staticmethod
    def _dict_to_text(data: Dict[str, Any], indent: int = 0) -> str:
        """
        Convert dictionary to formatted text string.
        Recursively handles nested dicts and lists.
        
        Args:
            data: Dictionary to convert
            indent: Indentation level for nested items
        
        Returns:
            str: Formatted text representation
        """
        lines = []
        for key, value in data.items():
            if value is None or value == '':
                continue
            
            # Format key
            key_str = key.replace('_', ' ').title()
            indent_str = "  " * indent
            
            if isinstance(value, dict):
                # Nested dict
                lines.append(f"\n{indent_str}{key_str}:")
                nested_text = DocumentExporter._dict_to_text(value, indent + 1)
                lines.append(nested_text)
            elif isinstance(value, list):
                # List of items
                lines.append(f"\n{indent_str}{key_str}:")
                for item in value:
                    if isinstance(item, dict):
                        item_text = DocumentExporter._dict_to_text(item, indent + 1)
                        lines.append(item_text)
                    else:
                        lines.append(f"{indent_str}  • {str(item)}")
            else:
                # Simple value
                lines.append(f"{indent_str}{key_str}: {str(value)}")
        
        return "\n".join(lines)

    # ========================
    # Utility Methods
    # ========================

    def generate_filename(
        self,
        org_name: str,
        tender_title: str
    ) -> str:
        """
        Generate a safe filename for the proposal.
        
        Args:
            org_name: Organization name
            tender_title: Tender title
        
        Returns:
            str: Safe filename with timestamp
        """
        # Sanitize names
        org_safe = "".join(c if c.isalnum() or c == " " else "" for c in org_name)[:20]
        tender_safe = "".join(c if c.isalnum() or c == " " else "" for c in tender_title)[:20]
        
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        
        filename = f"Proposal_{org_safe}_{tender_safe}_{timestamp}.docx"
        
        return filename.replace(" ", "_")


# Singleton instance
_document_exporter: Optional[DocumentExporter] = None


def get_document_exporter() -> DocumentExporter:
    """Get or create document exporter instance."""
    global _document_exporter
    if _document_exporter is None:
        _document_exporter = DocumentExporter()
    return _document_exporter
