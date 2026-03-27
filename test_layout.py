import sys
sys.path.insert(0, '.')

import logging
logging.disable(logging.CRITICAL)

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from io import BytesIO
from pathlib import Path

# Create a minimal test doc
doc = Document()

# Set margins
sections = doc.sections
for section in sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

# Add cover page placeholder
cover = doc.add_heading('PROPOSAL FOR TENDER', level=1)
cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('PROVISION OF BULK SMS SERVICES FOR A PERIOD OF THREE (3) YEARS')
doc.add_paragraph()
doc.add_page_break()

# Add org info placeholder
org_heading = doc.add_heading('Organization Information', level=2)
org_table = doc.add_table(rows=3, cols=2)
org_table.style = 'Light Grid Accent 1'
doc.add_page_break()

# Test heading with decorative line
heading_para = doc.add_heading('Executive Summary', level=1)
heading_para.paragraph_format.space_before = Pt(0)
heading_para.paragraph_format.space_after = Pt(12)

if heading_para.runs:
    for run in heading_para.runs:
        run.font.color.rgb = RGBColor(255, 99, 0)
        run.font.size = Pt(18)
        run.font.bold = True

# Add decorative line
line_para = doc.add_paragraph()
line_para.paragraph_format.space_before = Pt(0)
line_para.paragraph_format.space_after = Pt(12)
pPr = line_para._element.get_or_add_pPr()
pBdr = OxmlElement('w:pBdr')
bottom = OxmlElement('w:bottom')
bottom.set(qn('w:val'), 'single')
bottom.set(qn('w:sz'), '12')
bottom.set(qn('w:space'), '1')
bottom.set(qn('w:color'), 'FF6300')
pBdr.append(bottom)
pPr.append(pBdr)

# Add content
para = doc.add_paragraph('This is the executive summary with proper formatting and professional styling. Safaricom presents a comprehensive solution tailored to meet your requirements.')
para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
para.paragraph_format.space_before = Pt(6)
para.paragraph_format.space_after = Pt(6)

# Add page break
doc.add_page_break()

# Add another section
heading_para2 = doc.add_heading('Technical Approach', level=1)
heading_para2.paragraph_format.space_before = Pt(0)
heading_para2.paragraph_format.space_after = Pt(12)

if heading_para2.runs:
    for run in heading_para2.runs:
        run.font.color.rgb = RGBColor(255, 99, 0)
        run.font.size = Pt(18)
        run.font.bold = True

# Add decorative line
line_para2 = doc.add_paragraph()
line_para2.paragraph_format.space_before = Pt(0)
line_para2.paragraph_format.space_after = Pt(12)
pPr2 = line_para2._element.get_or_add_pPr()
pBdr2 = OxmlElement('w:pBdr')
bottom2 = OxmlElement('w:bottom')
bottom2.set(qn('w:val'), 'single')
bottom2.set(qn('w:sz'), '12')
bottom2.set(qn('w:space'), '1')
bottom2.set(qn('w:color'), 'FF6300')
pBdr2.append(bottom2)
pPr2.append(pBdr2)

# Add content
para2 = doc.add_paragraph('Our strategy involves establishing a high-performance SMS gateway with robust infrastructure capable of handling 1000 messages per second.')
para2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

# Save
output = BytesIO()
doc.save(output)
output.seek(0)
result = output.getvalue()

print(f'✅ Test document created successfully!')
print(f'✅ Document size: {len(result)} bytes')
print(f'✅ Sections with improved formatting:')
print(f'   - Cover page with proposal title')
print(f'   - Organization Information page')
print(f'   - Executive Summary with orange heading and decorative line')
print(f'   - Technical Approach with orange heading and decorative line')
print(f'   - Page breaks separating all sections')
print(f'   - Professional spacing and formatting')
